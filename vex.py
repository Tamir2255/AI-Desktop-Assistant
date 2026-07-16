from datetime import datetime
from pathlib import Path
import os
import math
import threading
import shutil
import asyncio
import tempfile
import json
import subprocess
import tkinter as tk
from tkinter import messagebox, simpledialog
import socket
import ctypes
from urllib.parse import quote_plus, urlparse, parse_qs, unquote
from urllib.request import Request, urlopen, build_opener, HTTPCookieProcessor
from urllib.error import URLError, HTTPError
import http.cookiejar
import webbrowser
import html
import re
import zipfile
import collections
from xml.sax.saxutils import escape

try:
    import pyperclip
except ImportError:
    pyperclip = None

# Suppress the RequestsDependencyWarning caused by library version mismatches
import warnings
warnings.filterwarnings("ignore", message=".*urllib3.*")
warnings.filterwarnings("ignore", message=".*chardet.*")

try:
    import pyttsx3
except ImportError:
    pyttsx3 = None

try:
    import edge_tts
except ImportError:
    edge_tts = None

try:
    import openai
except ImportError:
    openai = None

try:
    import google.generativeai as genai
except ImportError:
    genai = None

try:
    from elevenlabs import generate
except ImportError:
    generate = None

try:
    import speech_recognition as sr
except ImportError:
    sr = None

try:
    import yt_dlp
except ImportError:
    yt_dlp = None

try:
    import sympy
except ImportError:
    sympy = None

try:
    import matplotlib.pyplot as plt
    import numpy as np
except ImportError:
    plt = None
    np = None


APP_DIRECTORIES = [
    Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Windows" / "Start Menu" / "Programs",
    Path(os.environ.get("PROGRAMDATA", "")) / "Microsoft" / "Windows" / "Start Menu" / "Programs",
    Path(os.environ.get("LOCALAPPDATA", "")) / "Microsoft" / "WindowsApps",
    Path(os.environ.get("PROGRAMFILES", "")),
    Path(os.environ.get("PROGRAMFILES(X86)", "")),
]

BROWSER_NAMES = {
    "chrome": ["chrome", "google chrome"],
    "edge": ["edge", "microsoft edge", "msedge"],
    "firefox": ["firefox", "mozilla firefox"],
    "brave": ["brave", "brave browser"],
    "opera": ["opera", "opera browser"],
}

PYTHON_312 = Path("C:/Program Files/Python312/python.exe")
SPEAK_LISTENER = None
SPEAK_STATE_LISTENER = None
SEARCH_STATE_LISTENER = None

HOME_DIRECTORY = Path.home()
FOLDER_SEARCH_DIRECTORIES = [
    Path.cwd(),
    HOME_DIRECTORY / "Desktop",
    HOME_DIRECTORY / "Documents",
    HOME_DIRECTORY / "Downloads",
    HOME_DIRECTORY / "Pictures",
    HOME_DIRECTORY / "Music",
    HOME_DIRECTORY / "Videos",
]

PROTECTED_FOLDERS = {
    HOME_DIRECTORY.resolve(),
    Path.cwd().resolve(),
    Path(os.environ.get("WINDIR", "C:\\Windows")).resolve(),
    Path(os.environ.get("PROGRAMFILES", "C:\\Program Files")).resolve(),
    Path(os.environ.get("PROGRAMFILES(X86)", "C:\\Program Files (x86)")).resolve(),
}

DOCUMENT_OUTPUT_DIR = HOME_DIRECTORY / "Desktop" / "Vex Output"
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

# Global Cookie Handling to "approve" search engine terms
cookie_jar = http.cookiejar.CookieJar()
opener = build_opener(HTTPCookieProcessor(cookie_jar))

# Ensure output directory exists
DOCUMENT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

CONFIG_FILE = Path(r"c:\Users\User\Desktop\my projects\vex2\config.json")

def load_config():
    default_config = {
        "elevenlabs_api_key": "",
        "selected_voice": "Young Jamal",
        "theme_mode": "glass",
        "user_name": "",
        "last_run_date": "",
        "silent_startup": False,
        "user_birthday": "",
        "assistant_name": "Vex",
        "core_design": "Classic",
        "openai_api_key": "",
        "gemini_api_key": "",
        "target_language": "en"
    }
    if not CONFIG_FILE.exists():
        with open(CONFIG_FILE, 'w') as f:
            json.dump(default_config, f, indent=4)
        return default_config
    try:
        with open(CONFIG_FILE, 'r') as f:
            data = json.load(f)
            for key, value in default_config.items():
                data.setdefault(key, value)
            return data
    except:
        return default_config

def save_config(config_data):
    with open(CONFIG_FILE, 'w') as f:
        json.dump(config_data, f, indent=4)

CONFIG = load_config()
ELEVENLABS_API_KEY = CONFIG.get("elevenlabs_api_key", "")
SELECTED_VOICE = CONFIG.get("selected_voice", "Young Jamal")
USER_NAME = CONFIG.get("user_name", "")
ASSISTANT_NAME = CONFIG.get("assistant_name", "Vex")

HISTORY_FILE = Path(r"c:\Users\User\Desktop\my projects\vex2\vex_history.json")

def load_history():
    if HISTORY_FILE.exists():
        try:
            return json.loads(HISTORY_FILE.read_text())
        except:
            return []
    return []

def save_history(history):
    try:
        HISTORY_FILE.write_text(json.dumps(history, indent=4))
    except:
        pass


def get_current_time():
    return datetime.now().strftime("%I:%M %p")


def get_greeting(name=""):
    hour = datetime.now().hour
    greeting = "Good morning" if hour < 12 else "Good afternoon" if hour < 18 else "Good evening"
    return f"{greeting}, {name}" if name else greeting


# ---------------------------------------------------------------------------
# LOCAL / NO-EXTERNAL-BRAIN TEXT TOOLS
# ---------------------------------------------------------------------------
# Everything below works without any generative AI model (no OpenAI/Gemini
# calls). Summaries use plain word-frequency scoring, translation and grammar
# checks use small, free, keyless utility endpoints (not generative "brains"),
# and there is always a fully offline fallback.

STOPWORDS = frozenset("""
a about above after again against all am an and any are aren't as at be because
been before being below between both but by can't cannot could couldn't did
didn't do does doesn't doing don't down during each few for from further had
hadn't has hasn't have haven't having he he'd he'll he's her here here's hers
herself him himself his how how's i i'd i'll i'm i've if in into is isn't it
it's its itself let's me more most mustn't my myself no nor not of off on once
only or other ought our ours ourselves out over own same shan't she she'd
she'll she's should shouldn't so some such than that that's the their theirs
them themselves then there there's these they they'd they'll they're they've
this those through to too under until up very was wasn't we we'd we'll we're
we've were weren't what what's when when's where where's which while who
who's whom why why's with won't would wouldn't you you'd you'll you're you've
your yours yourself yourselves
""".split())

LANGUAGE_CODES = {
    "english": "en", "spanish": "es", "french": "fr", "german": "de",
    "italian": "it", "portuguese": "pt", "russian": "ru", "chinese": "zh",
    "mandarin": "zh", "japanese": "ja", "korean": "ko", "arabic": "ar",
    "hindi": "hi", "dutch": "nl", "turkish": "tr", "polish": "pl",
    "swedish": "sv", "greek": "el", "hebrew": "he", "vietnamese": "vi",
    "thai": "th", "indonesian": "id", "ukrainian": "uk", "swahili": "sw",
    "luganda": "lg", "romanian": "ro", "czech": "cs", "hungarian": "hu",
    "finnish": "fi", "danish": "da", "norwegian": "no",
}


def normalize_lang_code(name):
    name = (name or "en").strip().strip(".,!?").lower()
    return LANGUAGE_CODES.get(name, name[:5] if name else "en")


def local_summarize(text, num_sentences=3, as_bullets=False):
    """Pure frequency-based extractive summarizer. No external AI model used."""
    text = (text or "").strip()
    if not text:
        return "There is no text to summarize. Copy some text first, then try again."

    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return "There is no text to summarize."

    if len(sentences) <= num_sentences:
        summary = sentences
    else:
        words = re.findall(r"[a-zA-Z']+", text.lower())
        freq = collections.Counter(w for w in words if w not in STOPWORDS and len(w) > 2)

        if not freq:
            summary = sentences[:num_sentences]
        else:
            max_freq = max(freq.values())
            for w in freq:
                freq[w] = freq[w] / max_freq

            scored = []
            for i, sent in enumerate(sentences):
                sent_words = re.findall(r"[a-zA-Z']+", sent.lower())
                if not sent_words:
                    scored.append((0, i, sent))
                    continue
                score = sum(freq.get(w, 0) for w in sent_words) / len(sent_words)
                if i == 0:
                    score *= 1.2  # slight lead bias, common in extractive summarizers
                scored.append((score, i, sent))

            top = sorted(scored, key=lambda x: x[0], reverse=True)[:num_sentences]
            top_sorted = sorted(top, key=lambda x: x[1])
            summary = [s for _, _, s in top_sorted]

    if as_bullets:
        return "\n".join(f"\u2022 {s}" for s in summary)
    return " ".join(summary)


def translate_text(text, target_lang="en", source_lang="auto"):
    """Uses the free, keyless MyMemory translation endpoint. This is a plain
    machine-translation lookup service, not a generative AI 'brain'."""
    text = (text or "").strip()
    if not text:
        return "There is no text to translate. Copy some text first, then try again."

    target_lang = normalize_lang_code(target_lang)

    try:
        chunks = [text[i:i + 450] for i in range(0, len(text), 450)] or [text]
        translated_parts = []
        for chunk in chunks:
            langpair = f"{source_lang}|{target_lang}"
            url = f"https://api.mymemory.translated.net/get?q={quote_plus(chunk)}&langpair={quote_plus(langpair)}"
            req = Request(url, headers={"User-Agent": USER_AGENT})
            with opener.open(req, timeout=6) as response:
                payload = json.loads(response.read().decode("utf-8"))
                piece = payload.get("responseData", {}).get("translatedText", "").strip()
                translated_parts.append(piece if piece else chunk)
        return " ".join(translated_parts)
    except Exception:
        return "I couldn't reach the translation service. Please check your internet connection and try again."


def improve_text_basic(text):
    """Rule-based clarity pass: no generative AI involved, just tidy-up rules."""
    text = (text or "").strip()
    if not text:
        return "There is no text to improve."

    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\b(\w+)( \1\b)+', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\bi\b', 'I', text)
    text = re.sub(r'\s+([,.!?;:])', r'\1', text)

    sentences = re.split(r'(?<=[.!?])\s+', text)
    fixed = []
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        s = s[0].upper() + s[1:] if len(s) > 1 else s.upper()
        if s[-1] not in ".!?":
            s += "."
        fixed.append(s)
    return " ".join(fixed)


def check_grammar(text):
    """Uses the free, keyless LanguageTool public API (a rule/statistics based
    grammar checker, not a generative AI model) with an offline fallback."""
    text = (text or "").strip()
    if not text:
        return "There is no text to check. Copy some text first, then try again."

    try:
        url = "https://api.languagetool.org/v2/check"
        data = f"text={quote_plus(text)}&language=en-US".encode("utf-8")
        req = Request(url, data=data, headers={
            "User-Agent": USER_AGENT,
            "Content-Type": "application/x-www-form-urlencoded",
        })
        with opener.open(req, timeout=8) as response:
            payload = json.loads(response.read().decode("utf-8"))

        matches = payload.get("matches", [])
        if not matches:
            return f"No issues found. Your text looks good:\n\n{text}"

        corrected = text
        for m in sorted(matches, key=lambda x: x["offset"], reverse=True):
            replacements = m.get("replacements", [])
            if replacements:
                start, length = m["offset"], m["length"]
                corrected = corrected[:start] + replacements[0]["value"] + corrected[start + length:]

        notes = "\n".join(f"\u2022 {m.get('message', 'Possible issue')}" for m in matches[:5])
        return f"Corrected text:\n{corrected}\n\nNotes:\n{notes}"
    except Exception:
        return f"I couldn't reach the grammar service, so here's a basic cleanup instead:\n\n{improve_text_basic(text)}"


def ask_external_brain(query):
    """Optional: only used if the person has explicitly added their own
    OpenAI/Gemini key in settings. By default both keys are empty, so Vex
    runs with zero external AI models involved."""
    # Try OpenAI first
    oa_key = CONFIG.get("openai_api_key")
    if openai and oa_key:
        try:
            client = openai.OpenAI(api_key=oa_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": query}]
            )
            return response.choices[0].message.content
        except: pass
    
    # Fallback to Gemini
    gem_key = CONFIG.get("gemini_api_key")
    if genai and gem_key:
        try:
            genai.configure(api_key=gem_key)
            model = genai.GenerativeModel('gemini-pro')
            response = model.generate_content(query)
            return response.text
        except: pass
    
    return None


def is_connected():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        return False


def create_word_doc(filename, content, heading='Research Report'):
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(heading, 0)
        doc.add_paragraph(content)
        path = DOCUMENT_OUTPUT_DIR / f"{filename}.docx"
        doc.save(str(path))
        return path
    except ImportError:
        return "Error: python-docx not installed."


def create_excel_sheet(filename, data_list):
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Vex Export"
        for row in data_list:
            ws.append(row)
        path = DOCUMENT_OUTPUT_DIR / f"{filename}.xlsx"
        wb.save(str(path))
        return path
    except ImportError:
        return "Error: openpyxl not installed."


def create_pdf_doc(filename, content):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        path = DOCUMENT_OUTPUT_DIR / f"{filename}.pdf"
        c = canvas.Canvas(str(path), pagesize=letter)
        width, height = letter
        textobject = c.beginText(50, height - 50)
        textobject.setFont("Helvetica", 12)
        for line in content.split('\n'):
            textobject.textLine(line)
        c.drawText(textobject)
        c.showPage()
        c.save()
        return path
    except ImportError:
        return "Error: reportlab not installed."


def read_notifications():
    # Reading system notifications "without external brain" is highly OS specific.
    # On Windows, this usually requires the 'winrt' library or shell hooks.
    try:
        import win32com.client
        # This is a placeholder for actual shell notification access
        # Real implementation requires 'winrt' or complex 'ctypes'
        speak("Scanning system status...")
        return "System notifications are currently handled by Windows Action Center. I need 'winrt' installed to read them directly."
    except Exception:
        return "I cannot access notifications without the pywin32 library."


def get_weather(city=""):
    """Fetches weather using wttr.in scraping."""
    try:
        # Using format %C (Condition), %t (Temperature), %w (Wind)
        url = f"https://wttr.in/{quote_plus(city)}?format=%C+%t"
        # Force curl agent to get plain text; some systems need a simpler agent
        headers = {'User-Agent': 'curl/7.64.1'} 
        req = Request(url, headers=headers)
        with opener.open(req, timeout=5) as response:
            data = response.read().decode('utf-8').strip()
            # Ensure we didn't accidentally get an HTML page
            if "404" in data or "Unknown" in data or "<!DOCTYPE html>" in data:
                return "I couldn't find weather for that location."
            return f"The current weather in {city if city else 'your area'} is {data}."
    except:
        return "I'm having trouble connecting to the weather service."


def get_help_text():
    return (
        f"Here is what I can do. "
        "Open apps with open notepad or open chrome. "
        "Open browsers with open browser or open edge. "
        "Search the web with search, google, or look up. "
        "Use folders with open folder, where is folder, folder info, find folder, and delete folder. "
        "Use voice with voice, listen, or voice mode. "
        "Create documents with 'write in word', 'save as excel', or 'create pdf'. "
        "Ask me to 'write an essay on [topic]' for detailed research. "
        "Ask about the weather: 'weather in your current country' or 'how is the weather'. "
        "Use the Write tab for emails, rewrites, grammar fixes and translations. "
        "Use the Read tab to explain, summarize, translate, or get key points from any copied text. "
        "Use the Tools tab for the calculator, unit converter, graph plotter and more. "
        "Ask what time is it now, sum, add, hello, or what is your name. "
        f"Say bye to close {ASSISTANT_NAME}."
    )


def show_help():
    speak(get_help_text())


def normalize_name(name):
    return name.lower().replace(".lnk", "").replace(".exe", "").strip()


def set_interface_hooks(output_listener=None, speaking_listener=None, searching_listener=None):
    global SPEAK_LISTENER, SPEAK_STATE_LISTENER, SEARCH_STATE_LISTENER
    SPEAK_LISTENER = output_listener
    SPEAK_STATE_LISTENER = speaking_listener
    SEARCH_STATE_LISTENER = searching_listener


def emit_output(text):
    print(text) # Terminal fallback

    if SPEAK_LISTENER:
        SPEAK_LISTENER(text)


def set_speaking_state(is_speaking):
    if SPEAK_STATE_LISTENER:
        SPEAK_STATE_LISTENER(is_speaking)

def set_searching_state(is_searching):
    if SEARCH_STATE_LISTENER:
        SEARCH_STATE_LISTENER(is_searching)

async def _generate_neural_speech(text, voice, output_path):
    # Neural engines work best with clean text; they handle their own pauses.
    communicate = edge_tts.Communicate(text, voice, rate="+10%", pitch="-2Hz")
    await communicate.save(output_path)

def speak(text):
    emit_output(f"{ASSISTANT_NAME}: {text}")

    # Strip sources for cleaner speech
    speech_text = text.split("\n\nSources:")[0]

    set_speaking_state(True)
    winmm = ctypes.windll.winmm

    # --- Strategy 1: ElevenLabs (High-End Human Quality) ---
    if generate and ELEVENLABS_API_KEY and is_connected():
        try:
            audio_gen = generate(
                api_key=ELEVENLABS_API_KEY,
                text=speech_text,
                voice=SELECTED_VOICE,
                model="eleven_multilingual_v2"
            )
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                for chunk in audio_gen:
                    tmp_file.write(chunk)
                temp_path = tmp_file.name
            
            try:
                # Using short paths or proper quoting for MCI playback
                winmm.mciSendStringW(f'open "{temp_path}" type mpegvideo alias v_out', None, 0, 0)
                winmm.mciSendStringW(f'play v_out wait', None, 0, 0)
                winmm.mciSendStringW(f'close v_out', None, 0, 0)
            finally:
                if os.path.exists(temp_path): os.remove(temp_path)
                set_speaking_state(False)
                return
        except Exception as e:
            print(f"ElevenLabs failed: {e}")

    # --- Strategy 2: Microsoft Edge Neural (Best Free Strategy) ---
    if is_connected() and edge_tts:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                temp_path = tmp_file.name
            
            # AndrewNeural is much more human-like and expressive than Brian
            asyncio.run(_generate_neural_speech(speech_text, "en-US-AndrewNeural", temp_path))
            
            try:
                winmm.mciSendStringW(f'open "{temp_path}" type mpegvideo alias v_out', None, 0, 0)
                winmm.mciSendStringW('play v_out wait', None, 0, 0)
                winmm.mciSendStringW('close v_out', None, 0, 0)
            finally:
                if os.path.exists(temp_path): os.remove(temp_path)
                set_speaking_state(False)
            return
        except Exception as e:
            print(f"Neural TTS failed, falling back: {e}")

    # --- Strategy 3: Local Fallback (Only if Offline) ---
    if pyttsx3:
        try:
            engine = pyttsx3.init()
            engine.setProperty('rate', 170) 
            
            voices = engine.getProperty('voices')
            target_voices = [v for v in voices if "mark" in v.name.lower()] or [v for v in voices if "david" in v.name.lower()]
            if target_voices:
                engine.setProperty('voice', target_voices[0].id)

            engine.say(speech_text)
            engine.runAndWait()
        except Exception as e:
            emit_output(f"Voice output failed: {e}")
    
    set_speaking_state(False)


def respond(vex_text):
    speak(vex_text)


def listen_for_voice():
    if sr is None:
        speak("Voice input needs the SpeechRecognition package installed first.")
        emit_output("Install later with: pip install SpeechRecognition pyaudio")
        return ""

    recognizer = sr.Recognizer()

    try:
        with sr.Microphone() as source:
            speak("Listening.")
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            audio = recognizer.listen(source, timeout=5, phrase_time_limit=8)
    except AttributeError:
        speak("Voice input needs PyAudio installed for this Python version.")
        emit_output("PyAudio is installed in Python 3.12 on this PC, but not Python 3.14.")
        emit_output('Run Vex with: "C:/Program Files/Python312/python.exe" vex.py')
        emit_output("Or install Microsoft C++ Build Tools, then try installing PyAudio for Python 3.14 again.")
        return ""
    except OSError:
        speak("I could not access your microphone.")
        return ""
    except sr.WaitTimeoutError:
        speak("I did not hear anything.")
        return ""

    try:
        return recognizer.recognize_sphinx(audio).lower().strip()
    except sr.UnknownValueError:
        speak("I could not understand that.")
        return ""
    except sr.RequestError:
        pass

    try:
        speak("Offline voice recognition is not ready, so I will try online recognition.")
        return recognizer.recognize_google(audio).lower().strip()
    except sr.UnknownValueError:
        speak("I could not understand that.")
    except sr.RequestError:
        speak("Voice recognition could not connect. Check your internet or install PocketSphinx later.")

    return ""


def scan_apps():
    apps = {}

    for directory in APP_DIRECTORIES:
        if not directory.exists():
            continue

        for pattern in ("*.lnk", "*.exe"):
            for path in directory.rglob(pattern):
                app_name = normalize_name(path.stem)
                if app_name and app_name not in apps:
                    apps[app_name] = path

    for command in ("notepad", "calc", "mspaint", "cmd", "powershell"):
        command_path = shutil.which(command)
        if command_path:
            apps[command] = Path(command_path)

    return apps


def find_app(apps, app_name):
    app_name = normalize_name(app_name)

    if app_name in apps:
        return apps[app_name]

    for name, path in apps.items():
        if app_name in name or name in app_name:
            return path

    command_path = shutil.which(app_name)
    if command_path:
        return Path(command_path)

    return None


def open_app(apps, app_name):
    app_path = find_app(apps, app_name)

    if not app_path:
        speak(f"I could not find an app called {app_name}.")
        return

    try:
        os.startfile(app_path)
        speak(f"Opening {app_name}.")
    except OSError:
        try:
            subprocess.Popen([str(app_path)])
            speak(f"Opening {app_name}.")
        except OSError:
            speak(f"I found {app_name}, but Windows would not open it.")


def open_browser(apps, browser_name=""):
    browser_name = normalize_name(browser_name)

    if browser_name:
        open_app(apps, browser_name)
        return

    for browser_options in BROWSER_NAMES.values():
        for name in browser_options:
            browser = find_app(apps, name)
            if browser:
                open_app(apps, name)
                return

    webbrowser.open_new_tab("https://www.google.com")
    speak("Opening a new tab in your default browser.")


def scrape_site_content(url):
    """Attempts to fetch and extract a meaningful paragraph from a specific URL."""
    try:
        # Avoid non-html or large files
        if any(ext in url.lower() for ext in ['.pdf', '.jpg', '.zip', '.png']):
            return None
        headers = {
            'User-Agent': USER_AGENT,
            'Accept-Language': 'en-US,en;q=0.9'
        }
        req = Request(url, headers=headers)
        with opener.open(req, timeout=5) as response:
            raw_html = response.read().decode('utf-8', errors='ignore')
            # Look for text inside paragraph tags
            matches = re.findall(r'<(?:p|P)>(.*?)</(?:p|P)>|class="searchresult">(.*?)</div>', raw_html, re.DOTALL)
            results = []
            blacklist = ['login', 'register', 'sign up', 'password', 'user-agent', 'sitemap', 'allow:', 'disallow:', 'cookie', 'premium', 'unlock', 'access advanced', 'forgot your']
            
            for m in matches:
                p = m[0] if m[0] else m[1]
                clean_p = html.unescape(re.sub(r'<[^>]+>', '', p)).strip()
                # Remove Wikipedia-style citations [1], [a], [12][13] etc.
                clean_p = re.sub(r'\[\d+\]|\[[a-z]\]|\[edit\]', '', clean_p)
                clean_p = re.sub(r'\s+', ' ', clean_p) # Normalize spaces
                
                # Strict filtering: check length and ensure no boilerplate junk is present
                if len(clean_p) > 70 and not any(word in clean_p.lower() for word in blacklist):
                    results.append(clean_p)
            
            # Limit to top 8 paragraphs to prevent massive data dumps
            return "\n\n".join(results[:8]) if results else None
    except Exception:
        pass
    return None


def search_web(query, detailed=False, return_text=False):
    # Check external brain first for general questions
    if return_text and not detailed:
        brain_res = ask_external_brain(query)
        if brain_res: return brain_res

    query = query.strip()
    if not query:
        speak("Tell me what to search for.")
        return None

    # General search to allow DuckDuckGo AI/Snippets to lead the research
    search_query = query

    # Using DuckDuckGo HTML version for efficient scraping
    search_url = f"https://html.duckduckgo.com/html/?q={quote_plus(search_query)}"
    
    headers = {
        'User-Agent': USER_AGENT,
        'Accept-Language': 'en-US,en;q=0.9'
    }

    if not detailed and not return_text:
        webbrowser.open_new_tab(f"https://duckduckgo.com/?q={quote_plus(search_query)}")
        speak(f"Searching DuckDuckGo for {search_query}.")
        return None
    
    set_searching_state(True)
    answer = None
    sources = []
    try:
        req = Request(search_url, headers=headers)
        with opener.open(req) as response:
            html_content = response.read().decode('utf-8')

            # Extracting results from DuckDuckGo HTML structure
            raw_links = re.findall(r'class="result__a" href="([^"]+)"', html_content)
            titles = re.findall(r'class="result__a"[^>]*>(.*?)</a>', html_content, re.DOTALL)
            # Improved snippet extraction to handle different tags and variations
            snippets = re.findall(r'class="result__snippet[^"]*"[^>]*>(.*?)</(?:a|div)>', html_content, re.DOTALL)

            # Decode DuckDuckGo redirect links
            links = []
            for link in raw_links:
                if 'uddg=' in link:
                    try:
                        actual_url = unquote(link.split('uddg=')[1].split('&')[0])
                        links.append(actual_url)
                    except Exception:
                        links.append(link)
                else:
                    links.append(link)
            
            if return_text:
                clean_snippets = []
                for s in snippets:
                    txt = html.unescape(re.sub(r'<[^>]+>', '', s)).strip()
                    if len(txt) > 35 and not txt.startswith("http"):
                        clean_snippets.append(txt)
                
                # For Essays/Detailed research, aggregate snippets and perform deep analysis to form a full essay
                if detailed:
                    emit_output("Vex: Accessing AI sources for complete content...")
                    collected_text = []
                    
                    # Start with snippets as the "AI Overview"
                    if clean_snippets:
                        collected_text.append(" ".join(clean_snippets))
                    
                    # Visit links to "copy" the actual output from the target site
                    for i in range(min(len(links), 2)):
                        deep_txt = scrape_site_content(links[i])
                        if deep_txt and deep_txt not in collected_text:
                            collected_text.append(deep_txt)
                    
                    if collected_text:
                        answer = "\n\n".join(collected_text)
                
                if not answer and clean_snippets:
                    # Standard direct answer using top snippets
                    answer = "\n\n".join(clean_snippets[:2])

                for i in range(min(len(titles), 2)):
                    t = html.unescape(re.sub(r'<[^>]+>', '', titles[i]))
                    u = links[i] if i < len(links) else ""
                    if u: sources.append(f"{t}: {u}")

                # Wikipedia Fallback Logic
                if not answer and not detailed:
                    emit_output("Vex: DuckDuckGo direct answer not available. Checking Wikipedia...")
                    wiki_search_url = f"https://en.wikipedia.org/wiki/Special:Search?search={quote_plus(query)}"
                    wiki_res = scrape_site_content(wiki_search_url)
                    if wiki_res:
                        answer = wiki_res
                        sources.append(f"Wikipedia: {wiki_search_url}")
                
                # Deep Search Fallback: If answer is too short or missing, dive into top links
                if not answer and not detailed:
                    emit_output("Vex: Checking deeper sources...")
                    for i in range(min(len(links), 2)):
                        deep_answer = scrape_site_content(links[i])
                        if deep_answer:
                            answer = deep_answer
                            break

                if not answer:
                    answer = html.unescape(re.sub(r'<[^>]+>', '', titles[0])) if titles else "I couldn't find a direct answer."
                
                source_text = "\n\nSources:\n" + "\n".join(sources) if sources else ""
                return f"{answer}{source_text}"

            results = []
            for i in range(min(len(titles), 5)):
                clean_title = re.sub('<[^<]+?>', '', titles[i])
                clean_url = unquote(links[i].split('&')[0]) if i < len(links) else ""
                results.append(f"Source: {clean_title}\nLink: {clean_url}\n")
            
            return "\n".join(results)
    except Exception as e:
        emit_output(f"Search failed: {e}")
        webbrowser.open(search_url)
        answer = f"I couldn't scrape detailed info, but I opened the search page for {query}."
    finally:
        set_searching_state(False)
    return f"{answer}{source_text}" if 'source_text' in locals() else answer


def find_folder(folder_name):
    folder_name = folder_name.strip().strip('"')

    if not folder_name:
        return None

    possible_path = Path(folder_name).expanduser()
    if possible_path.exists() and possible_path.is_dir():
        return possible_path.resolve()

    normalized_folder_name = folder_name.lower()

    for directory in FOLDER_SEARCH_DIRECTORIES:
        if not directory.exists():
            continue

        try:
            for folder in directory.rglob("*"):
                if folder.is_dir() and folder.name.lower() == normalized_folder_name:
                    return folder.resolve()
        except (OSError, PermissionError):
            continue

    for directory in FOLDER_SEARCH_DIRECTORIES:
        if not directory.exists():
            continue

        try:
            for folder in directory.rglob("*"):
                if folder.is_dir() and normalized_folder_name in folder.name.lower():
                    return folder.resolve()
        except (OSError, PermissionError):
            continue

    return None


def open_folder(folder_name):
    folder_path = find_folder(folder_name)

    if not folder_path:
        speak(f"I could not find a folder called {folder_name}.")
        return

    try:
        os.startfile(folder_path)
        speak(f"Opening folder {folder_path.name}.")
        emit_output(f"Location: {folder_path}")
    except OSError:
        speak(f"I found the folder, but Windows would not open it.")
        emit_output(f"Location: {folder_path}")


def show_folder_info(folder_name):
    folder_path = find_folder(folder_name)

    if not folder_path:
        speak(f"I could not find a folder called {folder_name}.")
        return

    speak(f"The folder is located at {folder_path}.")
    emit_output(f"Name: {folder_path.name}")
    emit_output(f"Location: {folder_path}")


def is_safe_folder_to_delete(folder_path):
    resolved_path = folder_path.resolve()

    if resolved_path in PROTECTED_FOLDERS:
        return False

    if resolved_path.anchor == str(resolved_path):
        return False

    return True


def delete_folder(folder_name, confirmation_provider=None):
    folder_path = find_folder(folder_name)

    if not folder_path:
        speak(f"I could not find a folder called {folder_name}.")
        return

    if not is_safe_folder_to_delete(folder_path):
        speak("I will not delete that folder because it is protected or too important.")
        emit_output(f"Location: {folder_path}")
        return

    speak(f"I found the folder {folder_path.name}.")
    emit_output(f"Location: {folder_path}")

    if confirmation_provider:
        confirmed = confirmation_provider(folder_path)
    else:
        confirmation = input("Type DELETE to permanently delete this folder: ").strip()
        confirmed = confirmation == "DELETE"

    if not confirmed:
        speak("Folder deletion cancelled.")
        return

    try:
        shutil.rmtree(folder_path)
        speak(f"Deleted folder {folder_path.name}.")
    except OSError:
        speak("I tried to delete the folder, but Windows would not allow it.")


def understand_command(text):
    match = re.search(r"create a word document name it (.+?) and write an essay in it about (.+)", text)
    if match:
        return "write_essay_custom", (match.group(1).strip(), match.group(2).strip())

    search_file_match = re.search(r"search in (.+?) for (.+)", text)
    if search_file_match:
        return "search_in_file", (search_file_match.group(1).strip(), search_file_match.group(2).strip())

    find_text_match = re.search(r"find (.+?) in (.+)", text)
    if find_text_match:
        return "search_in_file", (find_text_match.group(2).strip(), find_text_match.group(1).strip())

    if "download" in text and ("youtube.com" in text or "youtu.be" in text):
        url_match = re.search(r'(https?://(?:www\.)?(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/shorts/)[^\s]+)', text)
        if url_match:
            return "download_video", url_match.group(0)

    graph_match = re.search(r"(?:graph|plot) (.+)", text)
    if graph_match:
        return "graph", graph_match.group(1).strip()

    conv_match = re.search(r"convert ([\d.]+) (\w+) to (\w+)", text)
    if conv_match:
        return "convert", (conv_match.group(1), conv_match.group(2), conv_match.group(3))

    if "weather" in text.lower() and not text.lower().startswith("search"):
        city = text.split("weather in")[-1].strip() if "weather in" in text else ""
        if "how is the weather" in text or text == "weather": city = ""
        return "weather", city

    if "shutdown" in text and ("pc" in text or "computer" in text):
        return "system_power", "shutdown"
    
    if "restart" in text and ("pc" in text or "computer" in text):
        return "system_power", "restart"

    if "lock" in text and ("pc" in text or "computer" in text or "windows" in text):
        return "system_power", "lock"

    if "volume up" in text:
        return "system_control", "volume_up"
    if "volume down" in text:
        return "system_control", "volume_down"
    if "mute" in text:
        return "system_control", "volume_mute"

    if "my name is" in text:
        return "set_name", text.split("my name is")[-1].strip()
    if "what is my name" in text:
        return "get_name", ""

    # READ/ANALYZE TAB INTENTS
    if "explain this highlighted text" in text.lower():
        # Split at the colon to get the actual content
        content = text.split("text:")[-1].strip() if "text:" in text else ""
        return "analyze_text", ("explain", content)

    # --- Local, no-external-brain text tools (Sider-style Read/Write actions) ---
    text_lower = text.lower()

    translate_match = re.search(r'translate\s+(?:this|the following)?\s*(?:text)?\s*to\s+([a-zA-Z]+)\s*[:\s]*(.*)', text, re.IGNORECASE | re.DOTALL)
    if translate_match:
        target = normalize_lang_code(translate_match.group(1))
        content = translate_match.group(2).strip()
        return "translate_text", (target, content)
    if text_lower.startswith("translate"):
        content = text.split(":", 1)[-1].strip() if ":" in text else re.sub(r'(?i)^translate(\s+this)?', '', text).strip()
        return "translate_text", (CONFIG.get("target_language", "en"), content)

    if re.search(r'\b(fix grammar|check grammar|grammar check)\b', text_lower):
        content = text.split(":", 1)[-1].strip() if ":" in text else ""
        return "grammar_check", content

    if re.search(r'\b(rewrite this|rewrite the following|improve this text|improve this)\b', text_lower):
        content = text.split(":", 1)[-1].strip() if ":" in text else ""
        return "rewrite_text", content

    if "key points" in text_lower:
        content = text.split(":", 1)[-1].strip() if ":" in text else ""
        return "key_points", content

    if "summarize" in text_lower:
        num_match = re.search(r'(\d+)\s*bullets?', text, re.IGNORECASE)
        num_sentences = int(num_match.group(1)) if num_match else 3
        if ":" in text:
            content = text.split(":", 1)[-1].strip()
        else:
            content = re.sub(r'(?i)summarize\s*(this|the following)?\s*(into\s*\d+\s*bullets?)?', '', text).strip()
        return "summarize_text", (num_sentences, content)

    if "clear chat" in text or "clear screen" in text:
        return "clear_chat", ""

    if any(phrase in text for phrase in ("help", "commands", "what can you do", "show help", "features", "capabilities")):
        return "help", ""

    if "write a professional email about" in text.lower() or "write an email about" in text.lower():
        # Extract the topic by splitting at 'about' and cleaning up colons/spaces
        topic = text.lower().split("about")[-1].strip().removeprefix(":").strip()
        return "write_email", topic

    if "write a social media post about" in text.lower() or "write a social post about" in text.lower():
        topic = text.lower().split("about")[-1].strip().removeprefix(":").strip()
        return "write_social_post", topic

    if "write an essay" in text or "research" in text:
        if "write an essay" in text:
            topic = text.split("essay")[-1].strip()
        else:
            topic = text.split("research")[-1].strip()
            
        # Clean up common connectors
        if topic.startswith("on "): topic = topic.removeprefix("on ").strip()
        if topic.startswith("about "): topic = topic.removeprefix("about ").strip()
        return "write_essay", topic

    if "write in word" in text or "create word" in text:
        return "create_doc", ("word", text)

    if "save as excel" in text or "create excel" in text:
        return "create_doc", ("excel", text)

    if "create pdf" in text or "write pdf" in text:
        return "create_doc", ("pdf", text)

    if "notifications" in text or "read alerts" in text:
        return "read_notifications", ""
        
    if text in ("voice mode", "start voice", "use voice"):
        return "start_voice_mode", ""

    if text in ("text mode", "stop voice", "use text"):
        return "start_text_mode", ""

    if text.startswith("zip folder "):
        return "zip_folder", text.removeprefix("zip folder ")

    if text.startswith("compress folder "):
        return "zip_folder", text.removeprefix("compress folder ")

    if text.startswith("open folder "):
        return "open_folder", text.removeprefix("open folder ")

    if text.startswith("delete folder "):
        return "delete_folder", text.removeprefix("delete folder ")

    if text.startswith("where is folder "):
        return "folder_info", text.removeprefix("where is folder ")

    if text.startswith("folder info "):
        return "folder_info", text.removeprefix("folder info ")

    if text.startswith("rename ") and " to " in text:
        parts = text.removeprefix("rename ").split(" to ")
        return "rename_item", (parts[0].strip(), parts[1].strip())

    if text.startswith("move ") and " to " in text:
        parts = text.removeprefix("move ").split(" to ")
        return "move_item", (parts[0].strip(), parts[1].strip())

    if text.startswith("delete file "):
        return "delete_file", text.removeprefix("delete file ")

    if text.startswith("find folder "):
        return "folder_info", text.removeprefix("find folder ")

    if text in ("new tab", "open new tab", "open a new tab"):
        return "open_browser", ""

    if text.startswith("open browser"):
        target = text.removeprefix("open browser").strip()
        if target.startswith("for "): target = target.removeprefix("for ").strip()
        if target: return "search_tab", target
        return "open_browser", ""

    if text.startswith("open web browser"):
        target = text.removeprefix("open web browser").strip()
        if target.startswith("for "): target = target.removeprefix("for ").strip()
        if target: return "search_tab", target
        return "open_browser", ""

    if text.startswith("search for "):
        return "search_tab", text.removeprefix("search for ")

    if text.startswith("new tab for "):
        return "search_tab", text.removeprefix("new tab for ")

    if text.startswith("open "):
        target = text.removeprefix("open ").strip()
        if "." in target or target.startswith("http") or target in ["youtube", "facebook", "google", "github", "gmail"]:
            return "open_url", target
        return "open_app", target

    if text.startswith("search "):
        return "web_search", text.removeprefix("search ")

    # Broader math detection for symbols like ^ or words like squared
    if any(op in text.lower() for op in ["+", "-", "*", "/", "^", "squared", "cubed"]) and any(c.isdigit() for c in text):
        math_expr = extract_math_expression(text.lower())
        if math_expr: return "calculate", math_expr

    # Catch questions that look like general knowledge
    if text.startswith(("what is", "who is", "how to", "why is", "tell me about")):
        # Only return web_search if it's not a specific internal query like "what is my name"
        if "my name" not in text and "time" not in text:
            return "web_search", text

    if text.startswith("google "):
        return "web_search", text.removeprefix("google ")

    if text.startswith("look up "):
        return "web_search", text.removeprefix("look up ")

    if text in ("voice", "listen", "voice input"):
        return "voice_input", ""

    return "chat", text


def find_path(name):
    name = name.strip().strip('"').lower()
    for directory in FOLDER_SEARCH_DIRECTORIES:
        if not directory.exists(): continue
        # Exact match
        for item in directory.iterdir():
            if item.name.lower() == name:
                return item.resolve()
        # Deep match
        try:
            for item in directory.rglob("*"):
                if item.name.lower() == name:
                    return item.resolve()
        except (OSError, PermissionError):
            continue
    return None


def rename_item(old_name, new_name):
    path = find_path(old_name)
    if not path:
        speak(f"I couldn't find {old_name}.")
        return
    try:
        new_path = path.parent / new_name
        os.rename(path, new_path)
        speak(f"Renamed {path.name} to {new_name}.")
    except Exception as e:
        speak(f"Failed to rename: {str(e)}")


def move_item(name, dest_folder_name):
    item_path = find_path(name)
    dest_path = find_folder(dest_folder_name)
    if not item_path or not dest_path:
        speak(f"I couldn't find the item or the destination.")
        return
    try:
        shutil.move(str(item_path), str(dest_path))
        speak(f"Moved {item_path.name} to {dest_path.name}.")
    except Exception as e:
        speak(f"Failed to move: {str(e)}")


def delete_file(file_name, confirmation_provider=None):
    path = find_path(file_name)
    if not path or not path.is_file():
        speak(f"I couldn't find a file called {file_name}.")
        return
    
    if not is_safe_folder_to_delete(path):
        speak("That file is protected.")
        return

    if confirmation_provider:
        confirmed = confirmation_provider(path)
    else:
        confirmed = input(f"Confirm deleting {path.name}? (y/n): ").lower() == 'y'

    if confirmed:
        try:
            os.remove(path)
            speak(f"Deleted file {path.name}.")
        except Exception as e:
            speak("Windows blocked the deletion.")


def zip_folder(folder_name):
    folder_path = find_folder(folder_name)
    if not folder_path:
        speak(f"I couldn't find a folder called {folder_name}.")
        return

    zip_name = f"{folder_path.name}_{datetime.now().strftime('%H%M%S')}.zip"
    zip_path = DOCUMENT_OUTPUT_DIR / zip_name

    try:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, folder_path)
                    zipf.write(file_path, arcname)
        speak(f"I have compressed {folder_path.name} into an archive at {zip_path}.")
    except Exception as e:
        speak(f"I failed to create the zip archive: {str(e)}")


def search_in_file(file_name, search_text):
    file_path = find_path(file_name)
    if not file_path or not file_path.is_file():
        speak(f"I couldn't find a file called {file_name}.")
        return

    try:
        found_lines = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for i, line in enumerate(f, 1):
                if search_text.lower() in line.lower():
                    found_lines.append(i)
        
        if found_lines:
            res = f"I found '{search_text}' in {file_path.name} on line(s): {', '.join(map(str, found_lines[:10]))}"
            if len(found_lines) > 10: res += " and more."
            speak(res)
        else:
            speak(f"I searched {file_path.name} but did not find the text '{search_text}'.")
    except Exception as e:
        speak(f"I couldn't read the file: {str(e)}")


def download_youtube_video(url):
    if yt_dlp is None:
        speak("I need the yt-dlp library to download videos.")
        emit_output("Please install it by running: pip install yt-dlp")
        return

    def run_download():
        try:
            speak("Starting your YouTube download. This may take a moment...")
            ydl_opts = {
                'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best',
                'outtmpl': str(DOCUMENT_OUTPUT_DIR / '%(title)s.%(ext)s'),
                'noplaylist': True,
            }
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
            speak("Download complete! I've saved the video to your Vex Output folder.")
        except Exception as e:
            speak(f"Sorry, the download failed: {str(e)}")

    threading.Thread(target=run_download, daemon=True).start()

def handle_graphing(expression):
    if sympy is None or plt is None:
        speak("I need sympy and matplotlib installed to draw graphs.")
        return

    try:
        # Clean y = x^2 into just x^2
        clean_expr = expression.lower().replace("y", "").replace("=", "").strip()
        expr = sympy.sympify(clean_expr)
        
        # Generate the plot
        from sympy.plotting import plot
        speak(f"Generating the graph for {expression}...")
        
        p = plot(expr, show=False, title=f"Graph of {expression}", line_color='#00F2FE')
        
        path = DOCUMENT_OUTPUT_DIR / f"Graph_{datetime.now().strftime('%H%M%S')}.png"
        p.save(str(path))
        
        os.startfile(path)
        respond(f"I've drawn the graph and saved it to your output folder.")
    except Exception as e:
        respond(f"I couldn't graph that. Error: {str(e)}")

def handle_unit_conversion(amount, u_from, u_to):
    # Mapping of common conversion factors relative to a base unit or direct formulas
    conversions = {
        'meters': {'feet': 3.28084, 'inches': 39.3701, 'kilometers': 0.001, 'miles': 0.000621371},
        'feet': {'meters': 0.3048, 'inches': 12, 'yards': 0.333333},
        'kilograms': {'pounds': 2.20462, 'grams': 1000},
        'pounds': {'kilograms': 0.453592, 'ounces': 16},
    }
    
    u_from, u_to = u_from.lower(), u_to.lower()
    try:
        val = float(amount)
        result = None
        
        # Temperature logic
        if u_from == "celsius" and u_to == "fahrenheit":
            result = (val * 9/5) + 32
            explanation = f"Step: Multiplied {val} by 1.8 and added 32.\nResult: {result}\u00b0F"
        elif u_from == "fahrenheit" and u_to == "celsius":
            result = (val - 32) * 5/9
            explanation = f"Step: Subtracted 32 from {val} and multiplied by 0.555.\nResult: {result}\u00b0C"
        # Standard lookup logic
        elif u_from in conversions and u_to in conversions[u_from]:
            factor = conversions[u_from][u_to]
            result = val * factor
            explanation = f"Step: Multiplied {val} {u_from} by the conversion factor {factor}.\nResult: {result} {u_to}"
        
        if result is not None:
            respond(explanation)
        else:
            respond(f"I don't have the conversion data for {u_from} to {u_to} yet.")
            
    except ValueError:
        respond("Please provide a valid number for conversion.")


def handle_math(expression):
    if sympy is None:
        speak("I need the SymPy library for my advanced calculator. Please run: pip install sympy")
        return

    try:
        # Clean expression for SymPy
        clean_expr = expression.replace("x", "*").replace("^", "**").strip()
        
        # Evaluate via SymPy
        parsed_expr = sympy.sympify(clean_expr)
        result = parsed_expr.evalf() if parsed_expr.is_number and not "nsimplify" in clean_expr else parsed_expr

        # Handle Fraction Conversion Specifically
        if "nsimplify" in clean_expr:
            explanation = f"Formula: {expression}\nMethod: Fraction Simplification\nTo convert the decimal, I identified the terminating place value and reduced the ratio to its simplest form.\nResult: {result}"
            respond(explanation)
            return

        # Handle Equation Solving with detailed method explanation
        if "=" in expression:
            parts = expression.split("=")
            lhs = sympy.sympify(parts[0].replace("solve", "").strip())
            rhs = sympy.sympify(parts[1].strip())
            eq = sympy.Eq(lhs, rhs)
            result = sympy.solve(eq)
            explanation = (
                f"Formula: {lhs} = {rhs}\n"
                f"Method: Algebraic Isolation\n"
                f"Step: Rewrote equality as standard symbolic equation.\n"
                f"Step: Isolate variables by performing inverse operations on both sides.\n"
                f"Result: {result}"
            )
        else:
            steps = [f"Formula: {parsed_expr}", "Method: Arithmetic / Exponentiation"]
            if parsed_expr.has(sympy.Pow): steps.append("- Exponents: Evaluated powers and roots first.")
            if parsed_expr.has(sympy.Mul): steps.append("- Multiplication/Division: Solved from left to right.")
            if parsed_expr.has(sympy.Add): steps.append("- Addition/Subtraction: Resolved remaining terms.")
            explanation = "Step-by-Step Breakdown:\n" + "\n".join(steps) + f"\nResult: {result}"

        respond(explanation)
    except Exception as e:
        respond("I couldn't solve that math problem. Please check the formatting or provide a valid mathematical expression.")


def extract_math_expression(text):
    """Extracts a mathematical expression from various natural language formats."""
    text = text.lower()
    
    # Handle fraction requests
    if any(phrase in text for phrase in ["fraction of", "as a fraction", "as fraction", "to fraction"]):
        num_match = re.search(r"(\d+\.?\d*)", text)
        if num_match: return f"nsimplify({num_match.group(1)})"

    # Word-to-Symbol operations
    text = text.replace("plus", "+").replace("minus", "-")
    text = text.replace("times", "*").replace("multiplied by", "*").replace("divided by", "/")

    # Handle word-based powers
    text = text.replace("squared", "**2").replace("cubed", "**3")
    text = re.sub(r"to the power of\s*(\d+)", r"**\1", text)
    text = text.replace("power of", "**").replace("power", "**")

    # Remove command keywords
    for word in ["calculate", "solve", "what is", "compute"]:
        text = text.replace(word, "").strip()
    
    # Clean up connectors like "add" or "sum of"
    for word in ["add", "sum of", "the "]:
        text = text.replace(word, "").strip()

    math_chars = "0123456789+-*/().^=x sqrt sin cos tan log nsimplify "
    cleaned = "".join([c for c in text if c in math_chars or c.isalpha()]).strip()
    
    if any(char.isdigit() for char in cleaned) or any(fn in cleaned for fn in ["sqrt", "sin", "nsimplify"]):
        return cleaned
    return None


class VexInterface:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Vex Sidebar")
        
        # Sidebar configuration
        self.sidebar_width = 380
        self.screen_width = self.root.winfo_screenwidth()
        self.screen_height = self.root.winfo_screenheight()
        
        # Floating Geometry: 85% height, centered vertically, 25px from right
        self.bar_height = int(self.screen_height * 0.85)
        self.y_pos = (self.screen_height - self.bar_height) // 2
        self.x_pos = self.screen_width - self.sidebar_width - 25
        
        self.root.geometry(f"{self.sidebar_width}x{self.bar_height}+{self.x_pos}+{self.y_pos}")
        self.root.configure(bg="#0A0D14", highlightthickness=0, bd=0)
        self.root.attributes("-topmost", True)  # Always on top
        self.root.attributes("-alpha", 0.85) # Initial transparency for glass look
        self.root.overrideredirect(True) # Remove standard title bar for extension look

        self.apps = {}
        self.voice_mode = False
        self.calc_mode = False
        self.showing_history = False
        self.current_tab = "chat" # chat, write, read, tools, history
        self.is_speaking = False
        self.is_searching = False
        self.is_collapsed = False
        self.pulse = 0
        self.action_buttons = []
        self.history = load_history()
        
        # Permanent Glass Theme Colors
        self.theme_colors = {
            "bg": "#070911",
            "header": "#10131C",
            "accent": "#22E5FF",
            "accent2": "#8B5CF6",
            "text": "#F0F6FC",
            "entry": "#161B26",
            "btn": "#1D2230",
            "btn_fg": "#22E5FF",
            "active_btn": "#2A3142",
            "vex_header": "#C792FF",
            "muted": "#8B95A6",
            "success": "#49FF9F",
            "ring_in": "#10131C", "ring1": "#22E5FF", "ring2": "#8B5CF6", "ring_out": "#161B26"
        }
        
        # Load Vex Learning data
        self.learning_data = {}
        learning_path = Path(r"c:\Users\User\Desktop\my projects\vex2\vex_learning.json")
        if learning_path.exists():
            try:
                self.learning_data = json.loads(learning_path.read_text())
            except: pass

        self.drag_data = {"x": 0, "y": 0}
        self.root.bind("<Button-1>", self.start_drag)
        self.root.bind("<B1-Motion>", self.do_drag)
        
        # Glass Transparency Toggling
        self.root.bind("<Enter>", lambda e: self.root.attributes("-alpha", 0.98))
        self.root.bind("<Leave>", lambda e: self.root.attributes("-alpha", 0.85))

        set_interface_hooks(self.add_output, self.set_speaking, self.set_searching)
        self.build_ui() # This will now call the new API settings dialog
        self.animate_core()
        
        self.root.update_idletasks()
        self.apply_rounded_corners(self.sidebar_width, self.bar_height, 40)

        threading.Thread(target=self.startup, daemon=True).start()

    def apply_rounded_corners(self, width, height, radius):
        """Applies smooth rounded corners to the window using Windows API."""
        try:
            hwnd = self.root.winfo_id()
            # Width and Height + 1 ensures the GDI region covers the entire anti-aliased edge
            region = ctypes.windll.gdi32.CreateRoundRectRgn(0, 0, width + 1, height + 1, radius, radius)
            ctypes.windll.user32.SetWindowRgn(hwnd, region, True)
        except Exception as e:
            print(f"Rounding failed: {e}")

    def build_ui(self):
        # Main Container
        self.main_container = tk.Frame(self.root, bg=self.theme_colors["bg"], bd=0, highlightthickness=0)
        self.main_container.pack(fill="both", expand=True)

        # Sider.ai Side Navigation Bar (Vertical strip for tabs)
        self.nav_bar = tk.Frame(self.main_container, bg=self.theme_colors["header"], width=50, bd=0, highlightthickness=0)
        self.nav_bar.pack(side="left", fill="y")

        def create_nav_btn(icon, tab_name):
            btn = tk.Button(self.nav_bar, text=icon, font=("Segoe UI", 14), 
                           bg=self.theme_colors["header"], fg=self.theme_colors["text"],
                           activebackground=self.theme_colors["accent"], relief="flat", bd=0,
                           command=lambda: self.switch_tab(tab_name))
            btn.pack(fill="x", pady=10)
            
            btn.bind("<Enter>", lambda e: btn.config(bg=self.theme_colors["btn"]) if self.current_tab != tab_name else None)
            btn.bind("<Leave>", lambda e: btn.config(bg=self.theme_colors["header"]) if self.current_tab != tab_name else None)
            return btn

        # Separator for the Nav Bar
        tk.Frame(self.nav_bar, bg=self.theme_colors["accent2"], height=2).pack(fill="x", padx=10, pady=5)

        self.nav_chat = create_nav_btn("\U0001F4AC", "chat")
        self.nav_write = create_nav_btn("\u270D\uFE0F", "write")
        self.nav_read = create_nav_btn("\U0001F4D6", "read")
        self.nav_tools = create_nav_btn("\U0001F6E0\uFE0F", "tools")
        self.nav_hist = create_nav_btn("\U0001F4DC", "history")

        # Body Container (Holds the content of the tabs)
        self.body_container = tk.Frame(self.main_container, bg=self.theme_colors["bg"], bd=0, highlightthickness=0)
        self.body_container.pack(side="right", fill="both", expand=True)

        # Custom Title Bar
        self.title_bar = tk.Frame(self.body_container, bg=self.theme_colors["header"], height=35, bd=0, highlightthickness=0)
        self.title_bar.pack(fill="x", side="top")

        self.title_label = tk.Label(self.title_bar, text=f"{ASSISTANT_NAME.upper()} ASSISTANT", fg=self.theme_colors["accent"], bg=self.theme_colors["header"], 
                 font=("Segoe UI", 9, "bold"))
        self.title_label.pack(side="left", padx=10)

        self.close_btn = tk.Button(self.title_bar, text="\u2715", command=self.root.destroy, bg=self.theme_colors["header"], fg="#ff4b4b",
                  relief="flat", font=("Arial", 10, "bold"), bd=0)
        self.close_btn.pack(side="right", padx=5)
        
        self.settings_btn = tk.Button(self.title_bar, text="\u22EE", command=self.show_settings_menu, bg=self.theme_colors["header"], fg=self.theme_colors["text"],
                  relief="flat", font=("Arial", 12, "bold"), bd=0)
        self.settings_btn.pack(side="right", padx=5)
        
        self.min_btn = tk.Button(self.title_bar, text="\u2014", command=self.toggle_sidebar, bg=self.theme_colors["header"], fg=self.theme_colors["text"],
                  relief="flat", font=("Arial", 10, "bold"), bd=0)
        self.min_btn.pack(side="right", padx=5)

        # Content
        self.content_frame = tk.Frame(self.body_container, bg=self.theme_colors["bg"], bd=0, highlightthickness=0)
        self.content_frame.pack(fill="both", expand=True, padx=10)

        self.canvas = tk.Canvas(
            self.content_frame,
            width=160,
            height=160,
            bg=self.theme_colors["bg"],
            highlightthickness=0,
        )
        self.canvas.pack(pady=(10, 5))

        self.status = tk.Label(
            self.content_frame,
            text=f"{ASSISTANT_NAME.upper()} ONLINE",
            fg=self.theme_colors["vex_header"],
            bg=self.theme_colors["bg"],
            font=("Segoe UI", 13, "bold"),
        )
        self.status.pack(pady=(0, 5))

        # --- INPUT AREA (Pack this FIRST with side="bottom" to ensure visibility) ---
        self.input_container = tk.Frame(self.content_frame, bg=self.theme_colors["bg"], bd=0, highlightthickness=0)
        self.input_container.pack(fill="x", side="bottom", pady=(0, 25)) # Increased pady for rounded edges

        # Styled Entry Frame (Box borders removed for a smoother design)
        self.entry_border = tk.Frame(
            self.input_container, 
            bg=self.theme_colors["accent2"], 
            highlightthickness=0, 
            bd=0 # Ensure no default border
        )
        self.entry_border.pack(fill="x", pady=(0, 12))

        self.entry = tk.Entry(
            self.entry_border,
            bg=self.theme_colors["header"],
            fg=self.theme_colors["text"],
            insertbackground=self.theme_colors["accent"],
            relief="flat",
            font=("Segoe UI", 11),
            bd=0,
            highlightthickness=0
        )
        self.entry.pack(fill="x", padx=2, pady=2, ipady=10)
        
        # Placeholder Logic
        self.placeholder = "Type your command here..."
        self.entry.insert(0, self.placeholder)
        self.entry.config(fg="gray")
        
        self.entry.bind("<FocusIn>", self._on_entry_click)
        self.entry.bind("<FocusOut>", self._on_focus_out)
        self.entry.bind("<Return>", lambda e: self.submit_command())

        # Action Button Row
        self.btn_row = tk.Frame(self.input_container, bg=self.theme_colors["bg"])
        self.btn_row.pack(fill="x")

        self.submit_btn = self.make_button(self.btn_row, "\u27A4", self.submit_command, bg=self.theme_colors["accent2"], fg=self.theme_colors["text"])
        self.submit_btn.pack(side="left", fill="x", expand=True, padx=2)
        self.voice_btn = self.make_button(self.btn_row, "\U0001F399", self.listen_once, bg=self.theme_colors["accent"], fg=self.theme_colors["bg"])
        self.voice_btn.pack(side="left", fill="x", expand=True, padx=2)
        self.stop_btn = self.make_button(self.btn_row, "\u23F9", self.stop_speaking, bg="#ff4b4b", fg="white")
        # stop_btn is not packed initially
        self.calc_btn = self.make_button(self.btn_row, "\U0001F9EE", self.toggle_calc_mode)
        self.calc_btn.pack(side="left", fill="x", expand=True, padx=2)
        self.make_button(self.btn_row, "\U0001F4CB", self.copy_last_answer).pack(side="left", fill="x", expand=True, padx=2)
        self.make_button(self.btn_row, "?", self.run_help).pack(side="left", padx=2)

        # --- OUTPUT AREA (Pack this SECOND to fill the remaining space) ---
        self.output_frame = tk.Frame(self.content_frame, bg=self.theme_colors["bg"])
        self.output_frame.pack(fill="both", expand=True, pady=(0, 10))
        self.output_frame.grid_columnconfigure(0, weight=1)
        self.output_frame.grid_rowconfigure(0, weight=1)

        self.scrollbar = tk.Scrollbar(
            self.output_frame,
            width=8,
            bg=self.theme_colors["btn"],
            troughcolor=self.theme_colors["bg"],
            activebackground=self.theme_colors["accent"],
            bd=0,
            highlightthickness=0,
            elementborderwidth=0
        )

        self.output = tk.Text(
            self.output_frame,
            bg=self.theme_colors["bg"],
            fg=self.theme_colors["text"],
            insertbackground=self.theme_colors["accent"],
            relief="flat",
            wrap="word",
            font=("Segoe UI", 10),
            padx=10,
            pady=10,
            yscrollcommand=self.scrollbar.set,
            bd=0,
            highlightthickness=0
        )
        self.output.grid(row=0, column=0, sticky="nsew")
        self.scrollbar.config(command=self.output.yview)

        # Hide/Show scrollbar on hover
        self.output_frame.bind("<Enter>", lambda e: self.scrollbar.grid(row=0, column=1, sticky="ns"))
        self.output_frame.bind("<Leave>", lambda e: self.scrollbar.grid_forget())

        # ChatGPT style tags
        self.output.tag_configure("user_header", foreground=self.theme_colors["accent2"], font=("Segoe UI", 9, "bold"), spacing1=15, justify="right")
        self.output.tag_configure("vex_header", foreground=self.theme_colors["vex_header"], font=("Segoe UI", 9, "bold"), spacing1=15)
        self.output.tag_configure("message_body", foreground=self.theme_colors["text"], spacing1=5, spacing3=5, lmargin1=15, lmargin2=15)
        self.output.tag_configure("user_body", foreground=self.theme_colors["text"], spacing1=5, spacing3=5, rmargin=15, justify="right")
        # New tags for calculator output
        self.output.tag_configure("math_formula", foreground=self.theme_colors["accent2"], font=("Segoe UI", 10, "italic"))
        self.output.tag_configure("math_method", foreground=self.theme_colors["accent"], font=("Segoe UI", 10, "bold"))
        self.output.tag_configure("math_step", foreground=self.theme_colors["text"], font=("Segoe UI", 9))
        self.output.tag_configure("bubble", background=self.theme_colors["header"])
        
        self.output.configure(state="disabled")

        # Floating Button (Hidden by default)
        self.float_btn = tk.Button(self.root, text="V", command=self.toggle_sidebar, 
                                  bg=self.theme_colors["accent"], fg="#031014",
                                  activebackground=self.theme_colors["accent"], 
                                  font=("Segoe UI", 14, "bold"), relief="flat", 
                                  bd=0, highlightthickness=0, borderwidth=0)

    # ------------------------------------------------------------------
    # Reusable scrollable panel + Sider-style action cards
    # ------------------------------------------------------------------
    def build_scroll_frame(self, parent):
        container = tk.Frame(parent, bg=self.theme_colors["bg"])
        canvas = tk.Canvas(container, bg=self.theme_colors["bg"], highlightthickness=0)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
        inner = tk.Frame(canvas, bg=self.theme_colors["bg"])

        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=inner, anchor="nw", width=self.sidebar_width - 70)
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        def _wheel(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")

        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _wheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))

        return container, inner

    def make_action_card(self, parent, icon, title, subtitle, command):
        card = tk.Frame(parent, bg=self.theme_colors["entry"], cursor="hand2")
        card.pack(fill="x", pady=6, padx=18)

        accent = tk.Frame(card, bg=self.theme_colors["accent"], width=4)
        accent.pack(side="left", fill="y")

        inner = tk.Frame(card, bg=self.theme_colors["entry"])
        inner.pack(side="left", fill="both", expand=True, padx=(12, 10), pady=12)

        top_row = tk.Frame(inner, bg=self.theme_colors["entry"])
        top_row.pack(fill="x", anchor="w")

        icon_lbl = tk.Label(top_row, text=icon, font=("Segoe UI", 13), bg=self.theme_colors["entry"], fg=self.theme_colors["accent"])
        icon_lbl.pack(side="left")
        title_lbl = tk.Label(top_row, text=title, font=("Segoe UI", 10, "bold"), bg=self.theme_colors["entry"], fg=self.theme_colors["text"])
        title_lbl.pack(side="left", padx=(8, 0))

        sub_lbl = tk.Label(inner, text=subtitle, font=("Segoe UI", 8), bg=self.theme_colors["entry"], fg=self.theme_colors["muted"],
                            anchor="w", justify="left", wraplength=self.sidebar_width - 120)
        sub_lbl.pack(fill="x", anchor="w", pady=(3, 0))

        widgets = [card, accent, inner, top_row, icon_lbl, title_lbl, sub_lbl]

        def on_click(event=None):
            command()

        def on_enter(event=None):
            for w in widgets:
                try:
                    w.config(bg=self.theme_colors["btn"])
                except tk.TclError:
                    pass

        def on_leave(event=None):
            for w in widgets:
                try:
                    w.config(bg=self.theme_colors["entry"])
                except tk.TclError:
                    pass

        for w in widgets:
            w.bind("<Button-1>", on_click)
            w.bind("<Enter>", on_enter)
            w.bind("<Leave>", on_leave)

        return card

    def switch_tab(self, tab_name):
        self.current_tab = tab_name
        # Reset visual states
        for btn in [self.nav_chat, self.nav_write, self.nav_read, self.nav_tools, self.nav_hist]:
            btn.config(bg=self.theme_colors["header"], fg=self.theme_colors["text"], relief="flat")
        
        # Hide everything by default
        self.canvas.pack_forget() 
        self.status.pack_forget()
        self.output_frame.pack_forget()
        self.input_container.pack_forget()
        self.hide_history_view()
        if hasattr(self, 'write_frame'): self.write_frame.pack_forget()
        if hasattr(self, 'read_frame'): self.read_frame.pack_forget()
        if hasattr(self, 'tools_frame'): self.tools_frame.pack_forget()

        if tab_name == "chat":
            self.nav_chat.config(fg=self.theme_colors["accent"], bg=self.theme_colors["btn"])
            self.canvas.pack(pady=(10, 5))
            self.status.pack(pady=(0, 5))
            self.input_container.pack(fill="x", side="bottom", pady=(0, 25))
            self.output_frame.pack(fill="both", expand=True, pady=(0, 10))
        elif tab_name == "write":
            self.nav_write.config(fg=self.theme_colors["accent"], bg=self.theme_colors["btn"])
            self.show_write_panel()
        elif tab_name == "read":
            self.nav_read.config(fg=self.theme_colors["accent"], bg=self.theme_colors["btn"])
            self.show_read_panel()
        elif tab_name == "tools":
            self.nav_tools.config(fg=self.theme_colors["accent"], bg=self.theme_colors["btn"])
            self.show_tools_panel()
        elif tab_name == "history":
            self.nav_hist.config(fg=self.theme_colors["accent"], bg=self.theme_colors["btn"])
            self.show_history_view()

    def show_write_panel(self):
        if not hasattr(self, 'write_frame'):
            self.write_frame, inner = self.build_scroll_frame(self.content_frame)
            tk.Label(inner, text="\u270D\uFE0F  WRITE ASSISTANT", fg=self.theme_colors["accent"], bg=self.theme_colors["bg"],
                     font=("Segoe UI", 13, "bold")).pack(pady=(20, 4), anchor="w", padx=18)
            tk.Label(inner, text="Pick a template, fill in the details, then hit send.", fg=self.theme_colors["muted"],
                     bg=self.theme_colors["bg"], font=("Segoe UI", 9), wraplength=self.sidebar_width - 70,
                     justify="left").pack(pady=(0, 14), anchor="w", padx=18)

            actions = [
                ("\U0001F4E7", "Professional Email", "Draft a polished email on any topic.", "Write a professional email about: "),
                ("\U0001F4AC", "Social Media Post", "Punchy copy for any platform.", "Write a social media post about: "),
                ("\U0001F9F9", "Rewrite & Improve", "Clean up wording, spacing and flow.", "Rewrite this: "),
                ("\u2705", "Fix Grammar", "Catch typos and grammar issues.", "Fix grammar: "),
                ("\U0001F310", "Translate My Text", "Translate typed text to any language.", "Translate this to english: "),
                ("\U0001F680", "Action Plan", "Turn a goal into concrete steps.", "Create an action plan for: "),
            ]
            for icon, title, subtitle, prompt in actions:
                self.make_action_card(inner, icon, title, subtitle, lambda p=prompt: self.apply_template(p))
        self.write_frame.pack(fill="both", expand=True)

    def show_read_panel(self):
        if not hasattr(self, 'read_frame'):
            self.read_frame, inner = self.build_scroll_frame(self.content_frame)
            tk.Label(inner, text="\U0001F4D6  READ & ANALYZE", fg=self.theme_colors["accent"], bg=self.theme_colors["bg"],
                     font=("Segoe UI", 13, "bold")).pack(pady=(20, 4), anchor="w", padx=18)
            tk.Label(inner, text="Highlight text anywhere, copy it (Ctrl+C), then tap an action.", fg=self.theme_colors["muted"],
                     bg=self.theme_colors["bg"], font=("Segoe UI", 9), wraplength=self.sidebar_width - 70,
                     justify="left").pack(pady=(0, 14), anchor="w", padx=18)

            actions = [
                ("\U0001F50D", "Explain Selection", "Get background and context.", "Explain this highlighted text: "),
                ("\U0001F4DD", "Summarize Selection", "Condense into a short summary.", "Summarize this into 3 bullets: "),
                ("\U0001F511", "Key Points", "Pull out the most important lines.", "Key points of: "),
                ("\U0001F310", "Translate Selection", "Translate the copied text.", "Translate this to english: "),
                ("\u2705", "Fix Grammar", "Clean up grammar and typos.", "Fix grammar: "),
                ("\U0001F4C4", "Analyze File Content", "Explain a file's content.", "Read and explain the file: "),
            ]
            for icon, title, subtitle, prompt in actions:
                self.make_action_card(inner, icon, title, subtitle, lambda p=prompt: self.analyze_clipboard(p))
        self.read_frame.pack(fill="both", expand=True)

    def show_tools_panel(self):
        if not hasattr(self, 'tools_frame'):
            self.tools_frame, inner = self.build_scroll_frame(self.content_frame)
            tk.Label(inner, text="\U0001F6E0\uFE0F  QUICK TOOLS", fg=self.theme_colors["accent"], bg=self.theme_colors["bg"],
                     font=("Segoe UI", 13, "bold")).pack(pady=(20, 4), anchor="w", padx=18)
            tk.Label(inner, text="Handy utilities that never need an internet 'brain'.", fg=self.theme_colors["muted"],
                     bg=self.theme_colors["bg"], font=("Segoe UI", 9), wraplength=self.sidebar_width - 70,
                     justify="left").pack(pady=(0, 14), anchor="w", padx=18)

            self.make_action_card(inner, "\U0001F9EE", "Calculator", "Toggle calculator mode for quick math.", self.quick_calc)
            self.make_action_card(inner, "\U0001F4D0", "Unit Converter", "Convert between units instantly.", self.quick_unit_convert)
            self.make_action_card(inner, "\U0001F4C8", "Graph Plotter", "Plot a function like y = x^2.", self.quick_graph)
            self.make_action_card(inner, "\u2600\uFE0F", "Weather Check", "Get the current weather anywhere.", self.quick_weather)
            self.make_action_card(inner, "\U0001F5DC\uFE0F", "Zip a Folder", "Compress a folder into a .zip file.", self.quick_zip)
            self.make_action_card(inner, "\U0001F50E", "Web Search", "Search the web for anything.", self.quick_search)
        self.tools_frame.pack(fill="both", expand=True)

    def quick_calc(self):
        self.switch_tab("chat")
        if not self.calc_mode:
            self.toggle_calc_mode()

    def quick_unit_convert(self):
        self.switch_tab("chat")
        self.entry.delete(0, tk.END)
        self.entry.config(fg=self.theme_colors["text"])
        self.entry.insert(0, "convert 10 meters to feet")
        self.entry.focus_set()

    def quick_graph(self):
        self.switch_tab("chat")
        self.entry.delete(0, tk.END)
        self.entry.config(fg=self.theme_colors["text"])
        self.entry.insert(0, "graph x^2")
        self.entry.focus_set()

    def quick_weather(self):
        self.switch_tab("chat")
        self.entry.delete(0, tk.END)
        self.entry.config(fg=self.theme_colors["text"])
        self.entry.insert(0, "weather in ")
        self.entry.focus_set()

    def quick_zip(self):
        self.switch_tab("chat")
        self.entry.delete(0, tk.END)
        self.entry.config(fg=self.theme_colors["text"])
        self.entry.insert(0, "zip folder ")
        self.entry.focus_set()

    def quick_search(self):
        self.switch_tab("chat")
        self.entry.delete(0, tk.END)
        self.entry.config(fg=self.theme_colors["text"])
        self.entry.insert(0, "search ")
        self.entry.focus_set()

    def _on_entry_click(self, event):
        if self.entry.get() in [self.placeholder, "Enter math problem..."]:
            self.entry.delete(0, "end")
            self.entry.insert(0, '')
            self.entry.config(fg=self.theme_colors["text"])

    def _on_focus_out(self, event):
        if self.entry.get() == '':
            self.entry.insert(0, self.placeholder)
            self.entry.config(fg="gray")

    def apply_template(self, prompt):
        self.switch_tab("chat")
        self.entry.delete(0, tk.END)
        self.entry.insert(0, prompt)
        self.entry.config(fg=self.theme_colors["text"])
        self.entry.focus_set()
        self.entry.icursor(tk.END)

    def analyze_clipboard(self, prompt):
        if pyperclip:
            text = pyperclip.paste()
            if text:
                self.switch_tab("chat")
                self.submit_command(f"{prompt} {text}")
            else:
                messagebox.showinfo("Read Mode", "Your clipboard is empty! Highlight and copy (Ctrl+C) some text first.")
        else:
            respond("Pyperclip library is missing. Use: pip install pyperclip")

    def toggle_history(self):
        self.switch_tab("history")

    def show_history_view(self):
        if not hasattr(self, 'history_view_frame'):
            self.history_view_frame = tk.Frame(self.content_frame, bg=self.theme_colors["bg"])
            tk.Label(self.history_view_frame, text="ACTIVITY HISTORY", fg=self.theme_colors["accent2"], 
                     bg=self.theme_colors["bg"], font=("Segoe UI", 10, "bold")).pack(pady=10)
            
            self.hist_canvas = tk.Canvas(self.history_view_frame, bg=self.theme_colors["bg"], highlightthickness=0)
            self.hist_scroll = tk.Scrollbar(self.history_view_frame, orient="vertical", command=self.hist_canvas.yview)
            self.hist_list_container = tk.Frame(self.hist_canvas, bg=self.theme_colors["bg"])
            
            self.hist_canvas.create_window((0, 0), window=self.hist_list_container, anchor="nw", width=self.sidebar_width-80)
            self.hist_canvas.configure(yscrollcommand=self.hist_scroll.set)
            
            self.hist_canvas.pack(side="left", fill="both", expand=True, padx=(10, 0))
            self.hist_scroll.pack(side="right", fill="y")
            
            self.hist_list_container.bind("<Configure>", lambda e: self.hist_canvas.configure(scrollregion=self.hist_canvas.bbox("all")))

        for widget in self.hist_list_container.winfo_children():
            widget.destroy()
            
        if not self.history:
            tk.Label(self.hist_list_container, text="No history yet", fg="gray", bg=self.theme_colors["bg"]).pack(pady=20)
        else:
            for item in reversed(self.history):
                f = tk.Frame(self.hist_list_container, bg=self.theme_colors["header"], pady=2)
                f.pack(fill="x", pady=2, padx=5)
                btn = tk.Button(f, text=item, command=lambda i=item: self.reuse_history_item(i),
                                bg=self.theme_colors["header"], fg=self.theme_colors["text"],
                                relief="flat", anchor="w", font=("Segoe UI", 9), wraplength=self.sidebar_width-80, justify="left")
                btn.pack(fill="x", side="left", expand=True)
        
        self.history_view_frame.place(x=0, y=0, relwidth=1, relheight=1)

    def hide_history_view(self):
        if hasattr(self, 'history_view_frame'):
            self.history_view_frame.place_forget()

    def reuse_history_item(self, text):
        self.switch_tab("chat")
        self.entry.delete(0, tk.END)
        self.entry.insert(0, text)
        self.submit_command()

    def toggle_sidebar(self):
        if not self.is_collapsed:
            # Shrink to a small button on the side
            self.main_container.pack_forget()
            self.root.geometry(f"60x60+{self.screen_width - 80}+100")
            self.apply_rounded_corners(60, 60, 60) # Perfect circle when collapsed
            self.float_btn.pack(fill="both", expand=True)
            self.is_collapsed = True
        else:
            # Expand back to full sidebar
            self.float_btn.pack_forget()
            self.root.geometry(f"{self.sidebar_width}x{self.bar_height}+{self.x_pos}+{self.y_pos}")
            self.apply_rounded_corners(self.sidebar_width, self.bar_height, 40)
            self.main_container.pack(fill="both", expand=True)
            self.is_collapsed = False

    def toggle_calc_mode(self):
        self.calc_mode = not self.calc_mode
        if self.calc_mode:
            self.calc_btn.config(fg="#ffffff", bg=self.theme_colors["accent2"])
            self.status.config(text="CALCULATOR MODE", fg=self.theme_colors["vex_header"])
            self.entry.delete(0, "end")
            self.entry.insert(0, "Enter math problem...")
            self.entry.config(fg="gray")
        else:
            self.calc_btn.config(fg=self.theme_colors["btn_fg"], bg=self.theme_colors["btn"])
            self.status.config(text="VEX ONLINE", fg=self.theme_colors["accent"])
            self.entry.delete(0, "end")
            self._on_focus_out(None)

    def show_settings_menu(self):
        menu = tk.Menu(self.root, tearoff=0, bg=self.theme_colors["header"], fg=self.theme_colors["text"], 
                       activebackground=self.theme_colors["accent"], activeforeground=self.theme_colors["bg"], font=("Segoe UI", 10))
        
        voice_menu = tk.Menu(menu, tearoff=0, bg=self.theme_colors["header"], fg=self.theme_colors["text"],
                             activebackground=self.theme_colors["accent"], activeforeground=self.theme_colors["bg"])
        voice_menu.add_command(label="Young Jamal", command=lambda: self.set_voice("Young Jamal"))
        voice_menu.add_command(label="Josh", command=lambda: self.set_voice("Josh"))
        menu.add_cascade(label="Select Voice", menu=voice_menu)
        
        menu.add_separator()
        menu.add_command(label="Personalization", command=self.open_name_settings)
        menu.add_command(label="API Settings", command=self.open_api_settings)
        menu.add_command(label="Clear History", command=self.clear_history)
        
        # Position the menu near the settings button
        x = self.root.winfo_rootx() + self.sidebar_width - 150
        y = self.root.winfo_rooty() + 35
        menu.post(x, y)

    def clear_history(self):
        if messagebox.askyesno("Clear History", "Permanently delete your activity history?"):
            self.history = []
            save_history([])
            if self.showing_history: self.show_history_view()
            speak("History cleared.")

    def open_api_settings(self):
        # Custom dialog for API key input
        dialog = tk.Toplevel(self.root)
        dialog.title("API Configuration")
        dialog.geometry("350x180")
        dialog.transient(self.root) # Make it appear on top of the main window
        dialog.grab_set() # Make it modal
        dialog.resizable(False, False)
        dialog.configure(bg=self.theme_colors["bg"])
        
        # Center the dialog
        self.root.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (dialog.winfo_width() // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")

        tk.Label(dialog, text="ElevenLabs API Key:", bg=self.theme_colors["bg"], fg=self.theme_colors["text"],
                 font=("Segoe UI", 10)).pack(pady=(20, 5))
        
        api_key_entry = tk.Entry(dialog, width=40, show='*', bg=self.theme_colors["entry"], fg=self.theme_colors["text"],
                                 insertbackground=self.theme_colors["accent"], relief="flat", bd=0, highlightthickness=0,
                                 font=("Segoe UI", 10))
        api_key_entry.insert(0, ELEVENLABS_API_KEY)
        api_key_entry.pack(pady=5)

        def save_key():
            global ELEVENLABS_API_KEY
            new_key = api_key_entry.get().strip()
            if new_key:
                ELEVENLABS_API_KEY = new_key
                CONFIG["elevenlabs_api_key"] = new_key
                save_config(CONFIG)
                speak("API Key updated successfully.")
            else:
                speak("API Key cannot be empty.")
            dialog.destroy()

        def cancel_dialog():
            dialog.destroy()

        btn_frame = tk.Frame(dialog, bg=self.theme_colors["bg"])
        btn_frame.pack(pady=15)

        tk.Button(btn_frame, text="Save", command=save_key, bg=self.theme_colors["accent"], fg="#031014",
                  activebackground=self.theme_colors["accent2"], relief="flat", font=("Segoe UI", 10, "bold"),
                  padx=15, pady=5, bd=0, highlightthickness=0).pack(side="left", padx=10)
        tk.Button(btn_frame, text="Cancel", command=cancel_dialog, bg=self.theme_colors["btn"], fg=self.theme_colors["text"],
                  activebackground=self.theme_colors["active_btn"], relief="flat", font=("Segoe UI", 10),
                  padx=15, pady=5, bd=0, highlightthickness=0).pack(side="left", padx=10)

        self.root.wait_window(dialog) # Wait for the dialog to close

    def open_name_settings(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Personalization")
        dialog.geometry("400x520")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.resizable(False, False)
        dialog.configure(bg=self.theme_colors["bg"])

        # Create a Canvas and Scrollbar
        canvas = tk.Canvas(dialog, bg=self.theme_colors["bg"], highlightthickness=0)
        scrollbar = tk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scroll_content = tk.Frame(canvas, bg=self.theme_colors["bg"])

        scroll_content.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scroll_content, anchor="nw", width=380)
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Mousewheel support
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        self.root.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (dialog.winfo_width() // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")

        # Name Setting
        tk.Label(scroll_content, text="What should Vex call you?", bg=self.theme_colors["bg"], fg=self.theme_colors["text"],
                 font=("Segoe UI", 10)).pack(pady=(20, 5))
        
        global USER_NAME
        name_entry = tk.Entry(scroll_content, width=30, bg=self.theme_colors["entry"], fg=self.theme_colors["text"],
                                 insertbackground=self.theme_colors["accent"], relief="flat", bd=0, highlightthickness=0,
                                 font=("Segoe UI", 10))
        name_entry.insert(0, USER_NAME)
        name_entry.pack(pady=5)

        tk.Label(scroll_content, text="What is my name?", bg=self.theme_colors["bg"], fg=self.theme_colors["text"], font=("Segoe UI", 10)).pack(pady=(10, 5))
        assistant_entry = tk.Entry(scroll_content, width=30, bg=self.theme_colors["entry"], fg=self.theme_colors["text"], relief="flat", font=("Segoe UI", 10))
        assistant_entry.insert(0, CONFIG.get("assistant_name", "Vex"))
        assistant_entry.pack(pady=5)

        tk.Label(scroll_content, text="Core Design:", bg=self.theme_colors["bg"], fg=self.theme_colors["text"], font=("Segoe UI", 10)).pack(pady=(10, 5))
        designs = ["Classic", "Pulsar", "Orbit", "Shield", "Wave"]
        design_var = tk.StringVar(value=CONFIG.get("core_design", "Classic"))
        design_menu = tk.OptionMenu(scroll_content, design_var, *designs)
        design_menu.config(bg=self.theme_colors["btn"], fg=self.theme_colors["text"], relief="flat")
        design_menu["menu"].config(bg=self.theme_colors["header"], fg=self.theme_colors["text"])
        design_menu.pack(pady=5)

        # Default Translate Language
        tk.Label(scroll_content, text="Default Translate Language:", bg=self.theme_colors["bg"], fg=self.theme_colors["text"],
                 font=("Segoe UI", 10)).pack(pady=(10, 5))
        lang_options = ["english", "spanish", "french", "german", "swahili", "luganda", "arabic", "chinese", "portuguese"]
        current_lang_name = next((k for k, v in LANGUAGE_CODES.items() if v == CONFIG.get("target_language", "en")), "english")
        lang_var = tk.StringVar(value=current_lang_name)
        lang_menu = tk.OptionMenu(scroll_content, lang_var, *lang_options)
        lang_menu.config(bg=self.theme_colors["btn"], fg=self.theme_colors["text"], relief="flat")
        lang_menu["menu"].config(bg=self.theme_colors["header"], fg=self.theme_colors["text"])
        lang_menu.pack(pady=5)

        # Birthday Setting
        tk.Label(scroll_content, text="Birthday (MM-DD):", bg=self.theme_colors["bg"], fg=self.theme_colors["text"],
                 font=("Segoe UI", 10)).pack(pady=(10, 5))
        
        bday_entry = tk.Entry(scroll_content, width=30, bg=self.theme_colors["entry"], fg=self.theme_colors["text"],
                                 insertbackground=self.theme_colors["accent"], relief="flat", bd=0, highlightthickness=0,
                                 font=("Segoe UI", 10))
        bday_entry.insert(0, CONFIG.get("user_birthday", ""))
        bday_entry.pack(pady=5)

        # Silent Startup Setting
        silent_var = tk.BooleanVar(value=CONFIG.get("silent_startup", False))
        tk.Checkbutton(scroll_content, text="Silent Startup", variable=silent_var, 
                       bg=self.theme_colors["bg"], fg=self.theme_colors["text"],
                       selectcolor=self.theme_colors["header"], activebackground=self.theme_colors["bg"],
                       font=("Segoe UI", 10), bd=0).pack(pady=10)

        def save_settings():
            canvas.unbind_all("<MouseWheel>")
            global USER_NAME, CONFIG, ASSISTANT_NAME
            new_name = name_entry.get().strip()
            new_bday = bday_entry.get().strip()
            new_assistant = assistant_entry.get().strip()
            
            if new_name:
                USER_NAME = new_name
                CONFIG["user_name"] = new_name
            if new_assistant:
                ASSISTANT_NAME = new_assistant
                CONFIG["assistant_name"] = new_assistant

            CONFIG["core_design"] = design_var.get()
            CONFIG["target_language"] = normalize_lang_code(lang_var.get())
            CONFIG["user_birthday"] = new_bday
            CONFIG["silent_startup"] = silent_var.get()
            save_config(CONFIG)
            
            if not CONFIG["silent_startup"]:
                speak(f"Personalization updated, {USER_NAME}.")
            else:
                emit_output(f"{ASSISTANT_NAME}: Personalization updated.")
            dialog.destroy()

        def cancel_dialog():
            canvas.unbind_all("<MouseWheel>")
            dialog.destroy()

        btn_frame = tk.Frame(scroll_content, bg=self.theme_colors["bg"])
        btn_frame.pack(pady=15)

        tk.Button(btn_frame, text="Save", command=save_settings, bg=self.theme_colors["accent"], fg="#031014",
                  activebackground=self.theme_colors["accent2"], relief="flat", font=("Segoe UI", 10, "bold"),
                  padx=15, pady=5, bd=0, highlightthickness=0).pack(side="left", padx=10)
        tk.Button(btn_frame, text="Cancel", command=cancel_dialog, bg=self.theme_colors["btn"], fg=self.theme_colors["text"],
                  activebackground=self.theme_colors["active_btn"], relief="flat", font=("Segoe UI", 10),
                  padx=15, pady=5, bd=0, highlightthickness=0).pack(side="left", padx=10)

    def set_voice(self, voice_name):
        global ELEVENLABS_API_KEY
        global SELECTED_VOICE
        SELECTED_VOICE = voice_name
        CONFIG["selected_voice"] = voice_name
        save_config(CONFIG)
        speak(f"Voice set to {voice_name}")

    def start_drag(self, event):
        # Ignore drag if clicking on interactive widgets to prevent accidental window movement
        if isinstance(event.widget, (tk.Button, tk.Scrollbar, tk.Text, tk.Entry)):
            self.drag_data["x"] = None
        else:
            self.drag_data["x"] = event.x
            self.drag_data["y"] = event.y

    def do_drag(self, event):
        if self.drag_data.get("x") is not None:
            x = self.root.winfo_x() - self.drag_data["x"] + event.x
            y = self.root.winfo_y() - self.drag_data["y"] + event.y
            self.root.geometry(f"+{x}+{y}")

    def make_button(self, parent, text, command, bg=None, fg=None):
        btn_bg = bg if bg else self.theme_colors["btn"]
        btn_fg = fg if fg else self.theme_colors["btn_fg"]
        active_bg = self.theme_colors["active_btn"] if not bg else bg

        btn = tk.Button(
            parent,
            text=text,
            command=command,
            bg=btn_bg,
            fg=btn_fg,
            activebackground=active_bg,
            activeforeground="#ffffff",
            relief="flat",
            bd=0,
            highlightthickness=0,
            font=("Segoe UI", 10, "bold"),
            padx=8,
            pady=6,
        )
        
        def on_enter(e):
            btn.config(bg=self.theme_colors["active_btn"] if not bg else bg)
        def on_leave(e):
            btn.config(bg=btn_bg, fg=btn_fg)
            
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)
        self.action_buttons.append(btn)
        return btn


    def add_output(self, text):
        self.root.after(0, self._add_output, text)

    def _add_output(self, text):
        self.output.configure(state="normal")
        start_idx = self.output.index("end-1c")

        # Determine if it's a math explanation and apply specific tags
        if "Formula:" in text and "Method:" in text and "Result:" in text:
            self.output.insert("end", f" \u25CF {ASSISTANT_NAME.upper()} \n", "vex_header")
            lines = text.split('\n')
            for line in lines:
                if line.startswith("Formula:"):
                    self.output.insert("end", f" {line}\n", "math_formula")
                elif line.startswith("Method:"):
                    self.output.insert("end", f" {line}\n", "math_method")
                elif line.startswith("Step:") or line.startswith("-"):
                    self.output.insert("end", f" {line}\n", "math_step")
                elif line.startswith("Result:"):
                    self.output.insert("end", f" {line}\n\n", "math_method") # Use method tag for result for emphasis
                else:
                    self.output.insert("end", f" {line}\n", "message_body")
        
        if text.startswith("You: "):
            msg = text.replace("You: ", "")
            self.output.insert("end", "YOU \u25CF \n", "user_header")
            self.output.insert("end", f"{msg} \n\n", ("user_body"))
        elif text.startswith("Vex: ") or text.startswith(f"{ASSISTANT_NAME}: "):
            prefix = f"{ASSISTANT_NAME}: "
            msg = text[len(prefix):] if text.startswith(prefix) else text.replace("Vex: ", "")
            self.output.insert("end", f" \u25CF {ASSISTANT_NAME.upper()} \n", "vex_header")
            self.output.insert("end", f" {msg} \n\n", "message_body") # Fallback for general Vex messages
        else:
            self.output.insert("end", f" {text} \n", "message_body")

        # Highlight links in the newly added block
        self._apply_link_tags(start_idx)

        self.output.see("end")
        self.output.configure(state="disabled")

    def clear_output_view(self):
        self.output.configure(state="normal")
        self.output.delete("1.0", tk.END)
        self.output.configure(state="disabled")

    def _apply_link_tags(self, start_pos):
        content = self.output.get(start_pos, "end")
        url_pattern = r'https?://[^\s\n\)\],]+'
        for match in re.finditer(url_pattern, content):
            tag_start = f"{start_pos} + {match.start()} chars"
            tag_end = f"{start_pos} + {match.end()} chars"
            self.output.tag_add("link", tag_start, tag_end)

    def _click_link(self, event):
        index = self.output.index(f"@{event.x},{event.y}")
        tags = self.output.tag_names(index)
        if "link" in tags:
            ranges = self.output.tag_ranges("link")
            for i in range(0, len(ranges), 2):
                if self.output.compare(ranges[i], "<=", index) and self.output.compare(index, "<=", ranges[i+1]):
                    url = self.output.get(ranges[i], ranges[i+1])
                    webbrowser.open(url)
                    break

    def set_speaking(self, is_speaking):
        self.root.after(0, self._set_speaking, is_speaking)

    def _set_speaking(self, is_speaking):
        self.is_speaking = is_speaking
        self.status.configure(text=f"{ASSISTANT_NAME.upper()} SPEAKING" if is_speaking else f"{ASSISTANT_NAME.upper()} ONLINE")
        if is_speaking:
            self.stop_btn.pack(side="left", fill="x", expand=True, padx=2, after=self.voice_btn)
        else:
            self.stop_btn.pack_forget()

    def stop_speaking(self):
        """Interrupts current MCI audio playback."""
        try:
            winmm = ctypes.windll.winmm
            winmm.mciSendStringW('stop v_out', None, 0, 0)
            winmm.mciSendStringW('close v_out', None, 0, 0)
        except Exception:
            pass
        self.set_speaking(False)

    def set_searching(self, state):
        self.root.after(0, self._set_searching, state)

    def _set_searching(self, state):
        self.is_searching = state

    def animate_core(self):
        self.canvas.delete("all")
        pulse_inc = 5.0 if self.is_searching else (3.0 if self.is_speaking else 1.0)
        self.pulse = (self.pulse + pulse_inc) % 60
        center = 100
        design = CONFIG.get("core_design", "Classic")
        
        main_color = "#FFD700" if self.is_searching else self.theme_colors["accent"]
        alt_color = "#FFFFFF" if self.is_searching else self.theme_colors["ring2"]

        if design == "Classic":
            base_radius = 35 + (self.pulse if self.is_speaking else (math.sin(self.pulse/10) * 5 + 5))
            rings = [(base_radius + 45, self.theme_colors["ring_out"]), (base_radius + 25, alt_color), (base_radius + 8, main_color)]
            for i, (radius, color) in enumerate(rings):
                self.canvas.create_oval(center-radius, center-radius, center+radius, center+radius, outline=color, width=2 if i==0 else 4)
        
        elif design == "Pulsar":
            # One thick glowing ring that breathes
            rad = 40 + math.sin(self.pulse/5) * 15
            self.canvas.create_oval(center-rad, center-rad, center+rad, center+rad, outline=main_color, width=8)
            self.canvas.create_oval(center-rad-5, center-rad-5, center+rad+5, center+rad+5, outline=alt_color, width=2)

        elif design == "Orbit":
            # Center circle with satellite dots
            self.canvas.create_oval(center-30, center-30, center+30, center+30, outline=main_color, width=3)
            for i in range(3):
                angle = (self.pulse * 5 + i * 120) * (math.pi / 180)
                ox, oy = center + 55 * math.cos(angle), center + 55 * math.sin(angle)
                self.canvas.create_oval(ox-5, oy-5, ox+5, oy+5, fill=alt_color, outline=main_color)

        elif design == "Shield":
            # Rotating Hexagon-like Diamond
            pts = []
            for i in range(4):
                angle = (self.pulse * 4 + i * 90) * (math.pi / 180)
                rad = 50 + (math.sin(self.pulse/5) * 10)
                pts.extend([center + rad * math.cos(angle), center + rad * math.sin(angle)])
            self.canvas.create_polygon(pts, outline=main_color, fill="", width=4)

        elif design == "Wave":
            # Horizontal Sine Wave across the core
            points = []
            for x in range(50, 150, 5):
                y = center + math.sin((x + self.pulse * 10) / 20) * 20
                points.extend([x, y])
            self.canvas.create_line(points, fill=main_color, smooth=True, width=3)
            self.canvas.create_oval(center-45, center-45, center+45, center+45, outline=alt_color, width=2)

        # Core Text
        core_rad = 30 + (self.pulse % 10 if self.is_speaking else 0)
        self.canvas.create_oval(center-core_rad, center-core_rad, center+core_rad, center+core_rad, fill=self.theme_colors["bg"], outline=main_color, width=2)
        self.canvas.create_text(center, center, text=ASSISTANT_NAME.upper(), fill=main_color if self.is_searching else self.theme_colors["text"], font=("Segoe UI", 12, "bold"))
        
        self.root.after(55, self.animate_core)

    def startup(self):
        global USER_NAME, CONFIG, ASSISTANT_NAME
        now = datetime.now()
        today = now.strftime("%Y-%m-%d")
        today_md = now.strftime("%m-%d")
        
        last_run = CONFIG.get("last_run_date", "")
        silent = CONFIG.get("silent_startup", False)
        birthday = CONFIG.get("user_birthday", "")
        
        greeting = get_greeting(USER_NAME)
        is_birthday = (birthday == today_md and birthday != "")

        if last_run != today:
            CONFIG["last_run_date"] = today
            save_config(CONFIG)
            msg = f"{greeting}! I'm {ASSISTANT_NAME}."
            if is_birthday: msg += f" And Happy Birthday, {USER_NAME if USER_NAME else 'Boss'}!"
            if not silent: speak(msg)
            else: emit_output(f"{ASSISTANT_NAME}: {msg}")
        else:
            msg = f"Happy Birthday again, {USER_NAME}!" if is_birthday else f"{greeting}!"
            if not silent: speak(msg)
            else: emit_output(f"{ASSISTANT_NAME}: {msg}")
            
        self.apps = scan_apps()

    def submit_command(self, manual_cmd=None):
        if manual_cmd:
            command = manual_cmd.strip()
        else:
            if self.entry.get() in [self.placeholder, "Enter math problem..."]:
                return
            command = self.entry.get().strip()
            self.entry.delete(0, "end")

        if not command:
            return

        if not self.history or self.history[-1] != command:
            self.history.append(command)
            save_history(self.history)

        self.add_output(f"You: {command}")
        
        if self.calc_mode:
            math_expr = extract_math_expression(command.lower())
            if math_expr:
                threading.Thread(target=handle_math, args=(math_expr,), daemon=True).start()
            else:
                threading.Thread(target=respond, args=("Calculator mode is active. Please provide a math problem or exit calculator mode to search the web.",), daemon=True).start()
            return
            
        threading.Thread(target=self.process_command, args=(command,), daemon=True).start()

    def listen_once(self):
        threading.Thread(target=self._listen_once, daemon=True).start()

    def _listen_once(self):
        command = listen_for_voice()
        if command:
            self.add_output(f"You: {command}")
            self.process_command(command)

    def run_help(self):
        threading.Thread(target=self.process_command, args=("help",), daemon=True).start()

    def copy_last_answer(self):
        content = self.output.get("1.0", tk.END)
        start_marker = f" \u25CF {ASSISTANT_NAME.upper()} "
        start_idx = content.rfind(start_marker)
        
        if start_idx != -1:
            # Extract content after the last Vex header
            sub = content[start_idx + len(start_marker):]
            # Ends at Sources or next message header
            end_markers = ["\n\nSources:", " \u25CF YOU ", f" \u25CF {ASSISTANT_NAME.upper()} "]
            end_idx = len(sub)
            
            for marker in end_markers:
                m_pos = sub.find(marker)
                if m_pos != -1 and m_pos < end_idx:
                    end_idx = m_pos
            
            final_answer = sub[:end_idx].strip()
            
            if final_answer:
                self.root.clipboard_clear()
                self.root.clipboard_append(final_answer)
                # Provide visual feedback via the status label
                original_text = self.status.cget("text")
                self.status.config(text="COPIED!", fg=self.theme_colors["accent"])
                self.root.after(1500, lambda: self.status.config(text=original_text, fg=self.theme_colors.get("success", "#49ff9f")))

    def confirm_delete_folder(self, folder_path):
        result = {"confirmed": False}
        event = threading.Event()

        def ask():
            result["confirmed"] = messagebox.askyesno(
                "Delete folder",
                f"Permanently delete this folder?\n\n{folder_path}",
            )
            event.set()

        self.root.after(0, ask)
        event.wait()
        return result["confirmed"]

    def process_command(self, user_input):
        # Preserve original for case-sensitive data, but lower for logic
        user_input_lower = user_input.lower()
        
        # Handle conversational greetings
        if user_input_lower in ["hello vex", "hello bro", "bro", "wassup", "sup", "yo", "hi vex", "hi", "hello"]:
                respond("Hello! Vex here. What's up? How can I help you today?")
                return

        # Handle capability questions
        if "can you" in user_input_lower or "do you know how to" in user_input_lower:
            if "download" in user_input:
                respond("I certainly can! Just give me a YouTube link and say 'download' to save the video to your desktop.")
                return
            if "math" in user_input or "calculate" in user_input or "solve" in user_input:
                respond("I'm excellent at math! I can solve equations, convert decimals to fractions, and even draw graphs.")
                return
            if "essay" in user_input or "research" in user_input:
                respond("Yes, I can research any topic and write an essay for you in Word, PDF, or Excel.")
                return
            if "folder" in user_input or "file" in user_input:
                respond("I can manage your files: opening, moving, renaming, zipping, or deleting folders.")
                return
            if "translate" in user_input:
                respond("Yes! Say 'translate this to spanish: your text' and I'll handle it, no account or key needed.")
                return
            if "summarize" in user_input or "summary" in user_input:
                respond("I can! Copy any text and use 'Summarize Selection' in the Read tab, or say 'summarize: your text'.")
                return
            if "grammar" in user_input or "rewrite" in user_input:
                respond("Absolutely, try the Write tab for 'Fix Grammar' or 'Rewrite & Improve'.")
                return

        responses = {
            "how are you": "I'm just thinking about my creator, so I'm doing great!",
            "what is your name": "I'm Vex, a simple AI assistant created by Tamir.",
            "bye": "Goodbye, have a great day! Ohh and don't forget to tell Tasnim hi.",
            "haha": "Why are you laughing? That wasn't funny.",
            "wow": "Yeah I know, I am the coolest AI ever made.",
        }
        
        # Check learned responses first
        if user_input in self.learning_data:
            respond(self.learning_data[user_input])
            return

        if user_input in ("bye", "exit", "quit"):
            respond(responses["bye"])
            self.root.after(600, self.root.destroy)
            return

        if user_input in ("sum", "add"):
            respond("Use the console version for interactive sums for now.")
            return

        if user_input == "what time is it now":
            respond(f"The current time is {get_current_time()}")
            return

        intent, data = understand_command(user_input)

        if intent == "help":
            show_help()
            return

        if intent == "analyze_text":
            mode, content = data
            if not content:
                respond("I didn't see any text to analyze. Make sure you highlight and copy something first!")
                return
            respond(f"Analyzing that text for you now...")
            # Perform a deep search to find context or an explanation for the snippet
            info = search_web(f"{mode} this: {content}", return_text=True, detailed=True)
            if info:
                respond(info.split("\n\nSources:")[0])
            return

        if intent == "translate_text":
            target_lang, content = data
            if not content:
                respond("Please provide the text you want translated (or highlight text and copy it first).")
                return
            respond(f"Translating to {target_lang}...")
            result = translate_text(content, target_lang=target_lang)
            respond(result)
            return

        if intent == "grammar_check":
            if not data:
                respond("Please provide some text to check, or copy some text first.")
                return
            respond("Checking grammar...")
            result = check_grammar(data)
            respond(result)
            return

        if intent == "rewrite_text":
            if not data:
                respond("Please provide the text you'd like me to rewrite.")
                return
            respond("Rewriting for clarity...")
            result = improve_text_basic(data)
            respond(result)
            return

        if intent == "key_points":
            if not data:
                respond("Please provide the text you'd like key points from.")
                return
            result = local_summarize(data, num_sentences=5, as_bullets=True)
            respond(f"Key points:\n{result}")
            return

        if intent == "summarize_text":
            num_sentences, content = data
            if not content:
                respond("Please provide text to summarize, or copy some text first.")
                return
            result = local_summarize(content, num_sentences=num_sentences, as_bullets=(num_sentences > 1))
            respond(result)
            return

        if intent == "write_social_post":
            topic = data
            if not topic:
                respond("Please tell me what the post should be about.")
                return
            respond(f"Drafting a social post about {topic}...")
            info = search_web(f"engaging social media post ideas about {topic}", return_text=True, detailed=True)
            if info:
                draft = info.split("\n\nSources:")[0]
                respond(local_summarize(draft, num_sentences=2))
            return

        if intent == "write_email":
            topic = data
            if not topic:
                respond("Please tell me what the email should be about.")
                return
            respond(f"Drafting a professional email about {topic}...")
            # Use detailed search to find suitable templates and phrasing
            info = search_web(f"professional email template and content for {topic}", return_text=True, detailed=True)
            if info:
                draft = info.split("\n\nSources:")[0]
                respond(f"Subject: {topic.title()}\n\n{draft}")
            return
            
        if intent == "write_essay_custom":
            name, topic = data
            if not topic:
                respond("Please provide a topic for the essay.")
                return
            # Specifically use detailed mode for essays to get AI-like depth
            info = search_web(topic, return_text=True, detailed=True)
            if info:
                essay_content = info.split("\n\nSources:")[0]
                path = create_word_doc(name, essay_content, heading=topic)
                respond(f"I researched {topic} and saved the essay as '{name}.docx' at {path}. Please open the file and check for errors because the essay may not be perfect.")
            return
            
        if intent == "write_essay":
            if not data:
                respond("What topic should I write the essay about? For example: 'write an essay on technology'.")
                return
            # Specifically use detailed mode for essays to get AI-like depth
            info = search_web(data, return_text=True, detailed=True)
            if info:
                essay_content = info.split("\n\nSources:")[0]
                path = create_word_doc(f"Essay_{data.replace(' ', '_')}", essay_content, heading=data)
                respond(f"I have gathered information and saved it to a Word document at {path}. Please open the file and check for errors because the essay may not be perfect.")
            return

        if intent == "create_doc":
            doc_type, raw_text = data
            content = raw_text.split("content")[-1].strip() if "content" in raw_text else "Empty Document"
            # Clean up the content by removing the command triggers
            content = raw_text
            for trigger in ["write in word", "create word", "save as excel", "create excel", "create pdf", "write pdf"]:
                if trigger in content:
                    content = content.split(trigger)[-1].strip()
                    if content.startswith("content"): content = content.replace("content", "", 1).strip()
                    break
            if not content or content == raw_text: content = "Empty Document"
            name = f"Vex_{doc_type}_{datetime.now().strftime('%H%M%S')}"
            
            if doc_type == "word":
                res = create_word_doc(name, content)
            elif doc_type == "pdf":
                res = create_pdf_doc(name, content)
            elif doc_type == "excel":
                # Simple conversion of content to list of lists for excel
                data_rows = [[item.strip()] for item in content.split(',')]
                res = create_excel_sheet(name, data_rows)
                
            if isinstance(res, Path):
                respond(f"Document created successfully: {res.name}")
                os.startfile(res.parent)
            else:
                respond(res)
            return

        if intent == "read_notifications":
            msg = read_notifications()
            respond(msg)
            return

        if intent == "rename_item":
            rename_item(data[0], data[1])
            return

        if intent == "move_item":
            move_item(data[0], data[1])
            return

        if intent == "delete_file":
            delete_file(data, self.confirm_delete_folder)
            return

        if intent == "open_url":
            url = data
            if not url.startswith("http"):
                if "." not in url: url = f"https://www.{url}.com"
                else: url = f"https://{url}"
            webbrowser.open_new_tab(url)
            respond(f"Opening {url} in a new tab.")
            return

        if intent == "search_tab":
            search_web(data, return_text=False)
            return

        if intent == "calculate":
            handle_math(data)
            return

        if intent == "graph":
            handle_graphing(data)
            return

        if intent == "convert":
            handle_unit_conversion(*data)
            return

        if intent == "zip_folder":
            zip_folder(data)
            return

        if intent == "set_name":
            global USER_NAME
            USER_NAME = data
            CONFIG["user_name"] = data
            save_config(CONFIG)
            respond(f"Nice to meet you, {data}! I'll remember that.")
            return

        if intent == "get_name":
            if USER_NAME:
                respond(f"Your name is {USER_NAME}. At least, that's what you told me!")
            else:
                respond("I don't know your name yet. You can tell me by saying 'my name is' followed by your name.")
            return

        if intent == "system_control":
            # VK Codes: Volume Mute (0xAD), Volume Down (0xAE), Volume Up (0xAF)
            if data == "volume_up":
                for _ in range(5):
                    ctypes.windll.user32.keybd_event(0xAF, 0, 0, 0)
                respond("Turning volume up.")
            elif data == "volume_down":
                for _ in range(5):
                    ctypes.windll.user32.keybd_event(0xAE, 0, 0, 0)
                respond("Turning volume down.")
            elif data == "volume_mute":
                ctypes.windll.user32.keybd_event(0xAD, 0, 0, 0)
                respond("Muting audio.")
            return

        if intent == "weather":
            respond(get_weather(data))
            return

        if intent == "system_power":
            if data == "lock":
                ctypes.windll.user32.LockWorkStation()
                respond("Locking your PC.")
            elif messagebox.askyesno("Confirm", f"Are you sure you want to {data} the computer?"):
                if data == "shutdown":
                    os.system("shutdown /s /t 1")
                elif data == "restart":
                    os.system("shutdown /r /t 1")
            return

        if intent == "clear_chat":
            self.clear_output_view()
            respond("Chat cleared.")
            return

        if intent == "search_in_file":
            search_in_file(data[0], data[1])
            return

        if intent == "download_video":
            download_youtube_video(data)
            return

        # Remaining intent logic ...

        if intent == "start_voice_mode":
            self.voice_mode = True
            respond("Voice mode is on.")
            self.listen_once()
            return

        if intent == "start_text_mode":
            self.voice_mode = False
            respond("Text mode is on.")
            return

        if intent == "voice_input":
            self.listen_once()
            return

        if intent == "open_folder":
            open_folder(data)
            return

        if intent == "delete_folder":
            delete_folder(data, self.confirm_delete_folder)
            return

        if intent == "folder_info":
            show_folder_info(data)
            return

        if intent == "open_app":
            open_app(self.apps, data)
            return

        if intent == "open_browser":
            open_browser(self.apps, data)
            return

        if intent == "web_search":
            res = search_web(data, return_text=True)
            if res:
                respond(res)
            return

        for key, response in responses.items():
            if user_input_lower == key:
                respond(response)
                return

        # Try web search as fallback for general queries
        res = search_web(user_input, return_text=True)
        if res:
            respond(res)
        else:
            respond("I didn't understand that. Type help to see what I can do.")

    def run(self):
        self.root.mainloop()


def create_shortcuts():
    try:
        import win32com.client
    except ImportError:
        return

    script_path = Path(__file__).resolve()
    python_path = PYTHON_312 if PYTHON_312.exists() else Path(os.sys.executable)
    shortcut_targets = [
        Path.home() / "Desktop" / "Vex.lnk",
        Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Windows" / "Start Menu" / "Programs" / "Vex.lnk",
    ]

    shell = win32com.client.Dispatch("WScript.Shell")
    for shortcut_path in shortcut_targets:
        if shortcut_path.exists():
            continue

        try:
            shortcut = shell.CreateShortcut(str(shortcut_path))
            shortcut.TargetPath = str(python_path)
            shortcut.Arguments = f'"{script_path}"'
            shortcut.WorkingDirectory = str(script_path.parent)
            shortcut.IconLocation = str(python_path)
            shortcut.Description = "Launch Vex assistant"
            shortcut.Save()
        except OSError:
            continue


def simple_assistant():
    global USER_NAME
    speak(f"{get_greeting(USER_NAME)}! I'm Vex.")
    apps = scan_apps()
    speak("Try: open notepad, search Python tutorials, open folder Downloads, or voice mode.")
    speak("Type bye to exit.")

    responses = {
        "hello": "Hello! How can I help you today?",
        "hi": "Hi there! What do you need?",
        "how are you": "I'm just thinking about my creator, so I'm doing great!",
        "what is your name": "I'm Vex, a simple AI assistant created by Tamir.",
        "what can you do": "I can open apps, open browsers, search the web, open folders, delete folders with confirmation, use voice commands, summarize or translate text, fix grammar, and rewrite text for clarity.",
        "bye": "Goodbye, have a great day! Ohh and don't forget to tell Tasnim hi.",
        "haha": "Why are you laughing? That wasn't funny.",
        "wow": "Yeah I know, I am the coolest AI ever made.",
    }

    def add_numbers(a, b):
        return a + b

    voice_mode = False

    while True:
        if voice_mode:
            user_input = listen_for_voice()
            if not user_input:
                continue
            speak(f"You said {user_input}.")
        else:
            user_input = input("You: ").lower().strip()

        command_text = user_input

        if user_input == "bye":
            respond(responses["bye"])
            break

        if user_input in ("sum", "add"):
            a = int(input("Enter 1st number: "))
            b = int(input("Enter 2nd number: "))
            respond(f"Sum of {a} and {b} is {add_numbers(a, b)}")
            continue

        if user_input == "what time is it now":
            respond(f"The current time is {get_current_time()}")
            continue

        intent, data = understand_command(user_input)

        if intent == "help":
            show_help()
            continue

        if intent == "translate_text":
            target_lang, content = data
            if content:
                respond(translate_text(content, target_lang=target_lang))
            else:
                respond("Please provide the text you want translated.")
            continue

        if intent == "grammar_check":
            if data:
                respond(check_grammar(data))
            else:
                respond("Please provide some text to check.")
            continue

        if intent == "rewrite_text":
            if data:
                respond(improve_text_basic(data))
            else:
                respond("Please provide the text you'd like rewritten.")
            continue

        if intent == "key_points":
            if data:
                respond("Key points:\n" + local_summarize(data, num_sentences=5, as_bullets=True))
            else:
                respond("Please provide the text you'd like key points from.")
            continue

        if intent == "summarize_text":
            num_sentences, content = data
            if content:
                respond(local_summarize(content, num_sentences=num_sentences, as_bullets=(num_sentences > 1)))
            else:
                respond("Please provide text to summarize.")
            continue

        if intent == "write_essay_custom":
            name, topic = data
            if not topic:
                respond("Please provide a topic.")
                continue
            info = search_web(topic, return_text=True, detailed=True)
            if info:
                essay_content = info.split("\n\nSources:")[0]
                path = create_word_doc(name, essay_content, heading=topic)
                respond(f"Document '{name}.docx' created with research on {topic} at {path}. Please open the file and check for errors because the essay may not be perfect.")
            continue

        if intent == "write_essay":
            if not data:
                respond("Please provide a topic.")
                continue
            info = search_web(data, return_text=True, detailed=True)
            if info:
                essay_content = info.split("\n\nSources:")[0]
                path = create_word_doc(f"Research_{data}", essay_content, heading=data)
                respond(f"I researched {data} and saved the findings to {path}. Please open the file and check for errors because the essay may not be perfect.")
            continue

        if intent == "create_doc":
            doc_type, raw_text = data
            name = f"Console_Export_{datetime.now().strftime('%H%M%S')}"
            content = raw_text
            for trigger in ["write in word", "create word", "save as excel", "create excel", "create pdf", "write pdf"]:
                if trigger in content:
                    content = content.split(trigger)[-1].strip()
                    break
            if not content or content == raw_text: content = "Empty Document"
            name = f"Vex_Export_{datetime.now().strftime('%H%M%S')}"
            if doc_type == "word": create_word_doc(name, raw_text)
            elif doc_type == "pdf": create_pdf_doc(name, raw_text)
            respond(f"File {name} created.")
            respond(f"Document created as {name} in your Output folder.")
            continue
            
        if intent == "read_notifications":
            respond(read_notifications())
            continue

        if intent == "rename_item":
            rename_item(data[0], data[1])
            continue

        if intent == "move_item":
            move_item(data[0], data[1])
            continue

        if intent == "delete_file":
            delete_file(data)
            continue

        if intent == "open_url":
            url = data
            if not url.startswith("http"):
                if "." not in url: url = f"https://www.{url}.com"
                else: url = f"https://{url}"
            webbrowser.open_new_tab(url)
            respond(f"Opening {url}")
            continue

        if intent == "search_tab":
            search_web(data, return_text=False)
            continue

        if intent == "calculate":
            handle_math(data)
            continue

        if intent == "zip_folder":
            zip_folder(data)
            continue

        if intent == "search_in_file":
            search_in_file(data[0], data[1])
            continue

        if intent == "start_voice_mode":
            voice_mode = True
            respond("Voice mode is on. Say text mode to go back to typing.")
            continue

        if intent == "start_text_mode":
            voice_mode = False
            respond("Text mode is on.")
            continue

        if intent == "voice_input":
            voice_command = listen_for_voice()
            if voice_command:
                speak(f"You said {voice_command}.")
                command_text = voice_command
                intent, data = understand_command(voice_command)
            else:
                continue

        if intent == "open_folder":
            open_folder(data)
            continue

        if intent == "delete_folder":
            delete_folder(data)
            continue

        if intent == "folder_info":
            show_folder_info(data)
            continue

        if intent == "open_app":
            open_app(apps, data)
            continue

        if intent == "open_browser":
            open_browser(apps, data)
            continue

        if intent == "web_search":
            res = search_web(data, return_text=True)
            if res:
                respond(res)
            continue

        matched = False
        for key, response in responses.items():
            if command_text == key:
                respond(response)
                matched = True
                break

        if matched:
            continue

        # Fallback to web search for console too
        res = search_web(command_text, return_text=True)
        if res:
            respond(res)
        else:
            respond("I didn't understand that. My creator will soon update me. Be alert!")


if __name__ == "__main__":
    create_shortcuts()
    VexInterface().run()
